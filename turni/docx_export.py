"""Generazione del documento Word (DOCX) con il layout grafico richiesto."""
from __future__ import annotations
import re
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False,
                   align: int = WD_ALIGN_PARAGRAPH.LEFT,
                   font_size: int = 10, font_name: str = "Aptos",
                   color: RGBColor | None = None) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    if color is not None:
        font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_row_min_height(row, height_cm: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    twips = str(int((height_cm / 2.54) * 1440))
    tr_height.set(qn("w:val"), twips)
    tr_height.set(qn("w:hRule"), "atLeast")


def _format_docx_table(table, column_widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(column_widths_cm):
            row.cells[idx].width = Cm(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "110")
                node.set(qn("w:type"), "dxa")


def _extract_week_dates_labels(week_label: str) -> tuple[str, str]:
    """Estrae le etichette DATA in formato breve usato nel DOCX finale."""
    label = (week_label or "").strip()
    compact = re.sub(r"\s*-\s*", "-", label)
    match = re.search(
        r"(?i)(\d{1,2})\s*([A-Za-z\u00C0-\u00FF]{3,})?\s*-\s*(\d{1,2})\s*([A-Za-z\u00C0-\u00FF]{3,})?",
        compact,
    )
    if match:
        d1, m1, d2, m2 = match.groups()
        if not m1 and m2:
            m1 = m2
        if not m2 and m1:
            m2 = m1
        left = d1.zfill(2) + (f" {m1[:3].title()}" if m1 else "")
        right = d2.zfill(2) + (f" {m2[:3].title()}" if m2 else "")
        return left, right

    nums = re.findall(r"\b(\d{1,2})\b", label)
    if len(nums) >= 2:
        return nums[0].zfill(2), nums[1].zfill(2)
    if len(nums) == 1:
        return nums[0].zfill(2), ""
    return (label or "\u2014", "\u2014")


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def build_turni_docx(path: str, anno: str, result_rows: list[dict],
                     operatori: list[str], counts: list[int],
                     status_str: str, diff_val: int, penalty_val: int,
                     *, template: dict[str, Any] | None = None) -> None:
    """Genera il DOCX con il layout grafico continuo.

    *template* opzionale: dict con chiavi title, title_color, subtitle_color,
    font_title, font_body per personalizzare l'aspetto del documento.
    """
    tpl = template or {}
    title_text     = tpl.get("title", "AUDIO/VIDEO MESSINA-GANZIRRI")
    title_color    = tpl.get("title_color", "4C79C5")
    subtitle_color = tpl.get("subtitle_color", "A8D033")
    font_title     = tpl.get("font_title", "Times New Roman")
    font_body      = tpl.get("font_body", "Arial")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(0.75)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.55)
    section.right_margin = Cm(0.55)

    page_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm

    base_style = doc.styles["Normal"]
    base_style.font.name = font_body
    base_style.font.size = Pt(10)

    # Raggruppa per sede se presente
    sites = list(dict.fromkeys(r.get("site", "") for r in result_rows))
    has_sites = len(sites) > 1 or (sites and sites[0])

    ordered_months = list(dict.fromkeys(row["month"] for row in result_rows))
    months_caption = " - ".join(ordered_months) if ordered_months else "Programmazione"

    # -- Tabella titolo ----------------------------------------------------
    title_table = doc.add_table(rows=1, cols=1)
    title_table.style = "Table Grid"
    title_table.autofit = False
    title_cell = title_table.rows[0].cells[0]
    title_cell.width = Cm(page_width_cm)
    _set_cell_shading(title_cell, title_color)
    _clear_cell(title_cell)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = title_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(1)
    r1 = p1.add_run(title_text)
    r1.bold = True
    r1.italic = True
    r1.font.name = font_title
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = title_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"{months_caption}   {anno}")
    r2.bold = True
    r2.font.name = font_title
    r2.font.size = Pt(14)
    r2.font.color.rgb = _hex_to_rgb(subtitle_color)
    _set_row_min_height(title_table.rows[0], 1.45)

    # -- Tabelle dati (una per sede o una sola) ----------------------------
    def _add_data_table(rows_subset: list[dict], site_label: str = "") -> None:
        if site_label:
            site_p = doc.add_paragraph()
            site_p.paragraph_format.space_before = Pt(8)
            site_p.paragraph_format.space_after = Pt(2)
            sr = site_p.add_run(site_label)
            sr.bold = True
            sr.font.name = font_title
            sr.font.size = Pt(13)
            sr.font.color.rgb = _hex_to_rgb(title_color)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _format_docx_table(table, [4.65, 7.0, 7.0])

        hdr = table.rows[0].cells
        headers = ["DATA", "VIDEO", "AUDIO"]
        for idx, value in enumerate(headers):
            _set_cell_text(hdr[idx], value, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14,
                           font_name=font_title,
                           color=RGBColor(0x2D, 0x3E, 0x56))
            _set_cell_shading(hdr[idx], "C7D5EC")
        _set_row_min_height(table.rows[0], 0.95)

        body_fill = "E9E9E9"
        for row_data in rows_subset:
            left_label, right_label = _extract_week_dates_labels(
                row_data.get("week", ""))

            giovedi = table.add_row().cells
            giovedi_values = [
                left_label or "\u2014",
                row_data.get("video", "") or "",
                row_data.get("audio", "") or "",
            ]
            for idx, value in enumerate(giovedi_values):
                _set_cell_text(
                    giovedi[idx], str(value),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    font_size=10, font_name=font_body)
                _set_cell_shading(giovedi[idx], body_fill)
            _set_row_min_height(table.rows[-1], 0.55)

            domenica = table.add_row().cells
            merged = domenica[1].merge(domenica[2])
            _set_cell_text(domenica[0], right_label or "\u2014",
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           font_size=10, font_name=font_body)
            _set_cell_shading(domenica[0], body_fill)
            _set_cell_text(merged, row_data.get("sabato", "") or "",
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           font_size=10, font_name=font_body)
            _set_cell_shading(merged, body_fill)
            _set_row_min_height(table.rows[-1], 0.55)

    if has_sites:
        for site in sites:
            site_rows = [r for r in result_rows if r.get("site", "") == site]
            _add_data_table(site_rows, site_label=site or "Sede principale")
    else:
        _add_data_table(result_rows)

    footer = section.footer.paragraphs[0]
    footer.clear()

    doc.save(path)
