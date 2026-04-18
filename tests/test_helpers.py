"""Test suite per turni v16 — importa dai moduli del package turni/."""
import json
import os
import time
import pathlib
import tempfile
import threading
import unittest
from unittest import mock

from turni.helpers import (
    normalize_name,
    _safe_csv_cell,
    _dedupe_normalized_texts,
    _group_weeks_by_normalized_month,
    _order_weeks_by_declared_months,
    _find_blank_text_entries,
    _normalized_week_key,
    _find_duplicate_week_keys,
)
from turni.validators import (
    SessionValidationError,
    _parse_step0_inputs,
    _validate_session_payload,
    _validate_solver_ready_weeks,
)
from turni.constants import LOCK_STALE_SECONDS
from turni.io_utils import (
    _write_text_file_atomic,
    _TargetFileLock,
    _is_stale_lock,
    _get_process_start_token,
    _pid_is_alive,
)
from turni.solver import TurniSolver, ORTOOLS_OK, SolvePhase
from turni.docx_export import build_turni_docx

try:
    from turni.ui.app import TurniApp
    _TKINTER_OK = True
except ImportError:
    _TKINTER_OK = False
    TurniApp = None  # type: ignore[assignment,misc]


class HelperTests(unittest.TestCase):
    def test_normalize_name_collapses_spaces_and_case(self):
        self.assertEqual(normalize_name("  Mario   Rossi  "), "mario rossi")

    def test_safe_csv_cell_prefixes_formula_like_values(self):
        self.assertEqual(_safe_csv_cell("=2+2"), "'=2+2")
        self.assertEqual(_safe_csv_cell("+1"), "'+1")
        self.assertEqual(_safe_csv_cell("-1"), "'-1")
        self.assertEqual(_safe_csv_cell(" @cmd"), "' @cmd")
        self.assertEqual(_safe_csv_cell("Mario Rossi"), "Mario Rossi")
        self.assertEqual(_safe_csv_cell(42), 42)

    def test_atomic_write_replaces_existing_file_contents(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = pathlib.Path(tmp) / "sessione.json"
            target.write_text("vecchio", encoding="utf-8")
            _write_text_file_atomic(str(target), '{"ok": true}')
            self.assertEqual(target.read_text(encoding="utf-8"), '{"ok": true}')

    def test_target_file_lock_blocks_second_writer(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = pathlib.Path(tmp) / "sessione.json"
            with _TargetFileLock(str(target)):
                with self.assertRaises(BlockingIOError):
                    with _TargetFileLock(str(target)):
                        self.fail("il secondo lock non dovrebbe essere acquisibile")
            self.assertFalse((pathlib.Path(str(target) + ".lock")).exists())

    def test_dedupe_normalized_texts_preserves_order_and_reports_duplicates(self):
        unique_values, unique_norms, duplicates = _dedupe_normalized_texts(
            ["Gennaio", " gennaio ", "Febbraio", "FEBBRAIO"]
        )
        self.assertEqual(unique_values, ["Gennaio", "Febbraio"])
        self.assertEqual(unique_norms, ["gennaio", "febbraio"])
        self.assertEqual(duplicates, [" gennaio ", "FEBBRAIO"])

    def test_group_weeks_by_normalized_month_uses_normalized_key(self):
        grouped = _group_weeks_by_normalized_month([
            {"month": "Gennaio", "week": "Sett.1"},
            {"month": " gennaio ", "week": "Sett.2"},
            {"month": "Febbraio", "week": "Sett.3"},
        ])
        self.assertEqual([w["week"] for w in grouped["gennaio"]], ["Sett.1", "Sett.2"])
        self.assertEqual([w["week"] for w in grouped["febbraio"]], ["Sett.3"])

    def test_order_weeks_by_declared_months_groups_months_without_interleaving(self):
        weeks = [
            {"month": "Gennaio", "week": "Sett. 1"},
            {"month": "Febbraio", "week": "Sett. 1"},
            {"month": "Marzo", "week": "Sett. 1"},
            {"month": "Gennaio", "week": "Sett. 2"},
            {"month": "Febbraio", "week": "Sett. 2"},
            {"month": "Marzo", "week": "Sett. 2"},
        ]
        ordered = _order_weeks_by_declared_months(
            weeks,
            ["Gennaio", "Febbraio", "Marzo"],
        )
        self.assertEqual(
            [(w["month"], w["week"]) for w in ordered],
            [
                ("Gennaio", "Sett. 1"),
                ("Gennaio", "Sett. 2"),
                ("Febbraio", "Sett. 1"),
                ("Febbraio", "Sett. 2"),
                ("Marzo", "Sett. 1"),
                ("Marzo", "Sett. 2"),
            ],
        )

    def test_find_blank_text_entries_detects_empty_or_space_only_strings(self):
        blanks = _find_blank_text_entries(["Mario", "   ", "", "Anna"])
        self.assertEqual(blanks, ["   ", ""])

    def test_normalized_week_key_ignores_case_and_extra_spaces(self):
        self.assertEqual(
            _normalized_week_key(" Gennaio ", "  Sett.1  "),
            ("gennaio", "sett.1"),
        )

    def test_find_duplicate_week_keys_uses_normalized_month_and_week(self):
        duplicates = _find_duplicate_week_keys([
            {"month": "Gennaio", "week": "Sett.1"},
            {"month": " gennaio ", "week": " sett.1 "},
            {"month": "Febbraio", "week": "Sett.1"},
        ])
        self.assertEqual(duplicates, [(" gennaio ", " sett.1 ")])

    def test_reset_result_state_clears_solver_state(self):
        solver = TurniSolver(["Mario"], [])
        solver.result_rows = [{"week": "Sett.1"}]
        solver.counts = [3]
        solver.diff_val = 2
        solver.penalty_val = 1
        solver.status_str = "OTTIMALE"
        solver.phase = SolvePhase.SOLVED   # sostituisce i 3 boolean flags
        solver.cp_solver = object()

        solver._reset_result_state()

        self.assertEqual(solver.result_rows, [])
        self.assertEqual(solver.counts, [])
        self.assertEqual(solver.diff_val, 0)
        self.assertEqual(solver.penalty_val, 0)
        self.assertEqual(solver.status_str, "")
        self.assertEqual(solver.phase, SolvePhase.IDLE)
        self.assertFalse(solver.solved_ok)
        self.assertFalse(solver.cancelled)
        self.assertFalse(solver.partial_result_available)
        self.assertIsNone(solver.cp_solver)

    @unittest.skipUnless(ORTOOLS_OK, "ortools non installato nel test environment")
    def test_solve_marks_cancelled_when_event_is_already_set(self):
        solver = TurniSolver(["Mario"], [])
        event = threading.Event()
        event.set()

        result = solver.solve(cancel_event=event)

        self.assertIn("Calcolo annullato", result)
        self.assertEqual(solver.phase, SolvePhase.CANCELLED)
        self.assertFalse(solver.solved_ok)
        self.assertFalse(solver.partial_result_available)
        self.assertTrue(solver.cancelled)

    @unittest.skipUnless(ORTOOLS_OK, "ortools non installato nel test environment")
    def test_solve_sets_error_phase_when_no_operators_available(self):
        # available=[] triggera il controllo early-return → phase ERROR
        solver = TurniSolver(["Mario", "Luca"], [
            {"month": "Gennaio", "week": "Sett.1",
             "available": [], "busy": ["mario", "luca"]},
        ])
        result = solver.solve()
        self.assertEqual(solver.phase, SolvePhase.ERROR)
        self.assertFalse(solver.solved_ok)
        self.assertFalse(solver.cancelled)
        self.assertIn("Errore", result)

    @unittest.skipUnless(ORTOOLS_OK, "ortools non installato nel test environment")
    def test_solve_sets_error_phase_when_infeasible(self):
        # Un solo operatore disponibile: ta != tv non soddisfacibile → INFEASIBLE
        solver = TurniSolver(["Mario", "Luca"], [
            {"month": "Gennaio", "week": "Sett.1",
             "available": [0], "busy": ["luca"]},
        ])
        result = solver.solve()
        self.assertEqual(solver.phase, SolvePhase.ERROR)
        self.assertFalse(solver.solved_ok)
        self.assertIn("Nessuna soluzione", result)

    def test_build_turni_docx_uses_requested_exact_layout(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = pathlib.Path(tmp) / "turni.docx"
            build_turni_docx(
                str(target),
                "2026",
                [
                    {
                        "month": "Gennaio",
                        "week": "Sett. 08-11 Gen",
                        "audio": "Ruben B.",
                        "video": "Riccardo N.",
                        "sabato": "Alessio A.",
                        "busy": "Maurizio B.",
                    },
                    {
                        "month": "Febbraio",
                        "week": "Sett. 01-05 Feb",
                        "audio": "Riccardo N.",
                        "video": "Ruben B.",
                        "sabato": "Maurizio B.",
                        "busy": "—",
                    },
                ],
                ["Ruben B.", "Riccardo N.", "Alessio A."],
                [2, 1, 1],
                "OTTIMALE",
                1,
                0,
            )
            self.assertTrue(target.exists())

            from docx import Document
            doc = Document(str(target))
            self.assertEqual(doc.sections[0].orientation, 1)
            self.assertFalse(any(par.text.strip() for par in doc.sections[0].footer.paragraphs))
            self.assertIn("AUDIO/VIDEO MESSINA-GANZIRRI", doc.tables[0].rows[0].cells[0].text)
            self.assertIn("Gennaio - Febbraio   2026", doc.tables[0].rows[0].cells[0].text)
            self.assertEqual(doc.tables[1].rows[0].cells[0].text, "DATA")
            self.assertEqual(doc.tables[1].rows[0].cells[1].text, "VIDEO")
            self.assertEqual(doc.tables[1].rows[0].cells[2].text, "AUDIO")
            self.assertEqual(doc.tables[1].rows[1].cells[0].text, "08 Gen")
            self.assertEqual(doc.tables[1].rows[1].cells[1].text, "Riccardo N.")
            self.assertEqual(doc.tables[1].rows[1].cells[2].text, "Ruben B.")
            self.assertEqual(doc.tables[1].rows[2].cells[0].text, "11 Gen")
            self.assertEqual(doc.tables[1].rows[2].cells[1].text, "Alessio A.")
            self.assertEqual(doc.tables[1].rows[2].cells[2].text, "Alessio A.")
            all_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertNotIn("RIEPILOGO TURNI PER OPERATORE", all_text)
            self.assertNotIn("DATI ELABORAZIONE", all_text)
            self.assertNotIn("NOTE / IMPEGNATI IN ALTRI REPARTI", all_text)


    def test_parse_step0_inputs_requires_at_least_two_operatori(self):
        with self.assertRaises(SessionValidationError):
            _parse_step0_inputs(
                "2026",
                "Gennaio",
                "Mario",
                "60",
            )

    def test_validate_solver_ready_weeks_rejects_less_than_two_available(self):
        with self.assertRaises(SessionValidationError):
            _validate_solver_ready_weeks(
                [{"month": "Gennaio", "week": "Sett.1", "busy_indices": [1]}],
                operator_count=2,
                error_prefix="Test",
            )


class SessionValidationTests(unittest.TestCase):
    def test_validate_session_payload_rejects_empty_operatori(self):
        with self.assertRaises(SessionValidationError):
            _validate_session_payload({
                "anno": "2026",
                "operatori": [],
                "mesi": ["Gennaio"],
                "weeks": [],
            })

    def test_validate_session_payload_rejects_empty_mesi(self):
        with self.assertRaises(SessionValidationError):
            _validate_session_payload({
                "anno": "2026",
                "operatori": ["Mario"],
                "mesi": [],
                "weeks": [],
            })

    def test_validate_session_payload_canonicalizes_strings_and_weeks(self):
        payload = _validate_session_payload({
            "anno": " 2026 ",
            "operatori": [" Mario Rossi ", "Anna Verdi", "Luca Neri"],
            "mesi": [" Gennaio "],
            "weeks": [{"month": " gennaio ", "week": " Sett.1 ", "busy_indices": [1]}],
            "solver_timeout": " 90 ",
        })
        self.assertEqual(payload["anno"], "2026")
        self.assertEqual(payload["operatori"], ["Mario Rossi", "Anna Verdi", "Luca Neri"])
        self.assertEqual(payload["mesi"], ["Gennaio"])
        self.assertEqual(payload["weeks"], [{"month": "gennaio", "week": "Sett.1", "busy_indices": [1]}])
        self.assertEqual(payload["solver_timeout"], "90")

    def test_validate_session_payload_rejects_duplicate_weeks(self):
        with self.assertRaises(SessionValidationError):
            _validate_session_payload({
                "anno": "2026",
                "operatori": ["Mario", "Anna"],
                "mesi": ["Gennaio"],
                "weeks": [
                    {"month": "Gennaio", "week": "Sett.1", "busy_indices": []},
                    {"month": " gennaio ", "week": " sett.1 ", "busy_indices": []},
                ],
            })

    def test_target_file_lock_recovers_stale_lock_with_dead_pid(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = pathlib.Path(tmp) / "sessione.json"
            lock_path = pathlib.Path(str(target) + ".lock")
            lock_path.write_text('{"pid": 999999, "created_at": 0}', encoding="utf-8")

            with _TargetFileLock(str(target)):
                self.assertTrue(lock_path.exists())

            self.assertFalse(lock_path.exists())

    def test_is_stale_lock_falls_back_to_old_mtime_for_malformed_metadata(self):
        with tempfile.TemporaryDirectory() as tmp:
            lock_path = pathlib.Path(tmp) / "sessione.json.lock"
            lock_path.write_text('not-json', encoding="utf-8")
            old = time.time() - (LOCK_STALE_SECONDS + 5)
            os.utime(lock_path, (old, old))
            self.assertTrue(_is_stale_lock(str(lock_path)))

    def test_validate_session_payload_rejects_single_operatore(self):
        with self.assertRaises(SessionValidationError):
            _validate_session_payload({
                "anno": "2026",
                "operatori": ["Mario"],
                "mesi": ["Gennaio"],
                "weeks": [],
            })

    def test_is_stale_lock_does_not_steal_live_pid_only_for_old_age(self):
        with tempfile.TemporaryDirectory() as tmp:
            lock_path = pathlib.Path(tmp) / "sessione.json.lock"
            lock_path.write_text(json.dumps({
                "pid": os.getpid(),
                "created_at": 0,
                "proc_start_token": _get_process_start_token(os.getpid()),
            }), encoding="utf-8")
            self.assertFalse(_is_stale_lock(str(lock_path), stale_after_seconds=1))

    def test_is_stale_lock_detects_pid_reuse_when_process_token_differs(self):
        with tempfile.TemporaryDirectory() as tmp:
            lock_path = pathlib.Path(tmp) / "sessione.json.lock"
            lock_path.write_text(json.dumps({
                "pid": os.getpid(),
                "created_at": time.time(),
                "proc_start_token": "old-token",
            }), encoding="utf-8")
            import turni.io_utils as io_utils_module
            with mock.patch.object(io_utils_module, "_pid_is_alive", return_value=True), \
                 mock.patch.object(io_utils_module, "_get_process_start_token", return_value="new-token"):
                self.assertTrue(_is_stale_lock(str(lock_path)))


class RestartAndTimeoutTests(unittest.TestCase):
    def test_validate_session_payload_rejects_invalid_solver_timeout(self):
        with self.assertRaises(SessionValidationError):
            _validate_session_payload({
                "anno": "2026",
                "operatori": ["Mario", "Anna"],
                "mesi": ["Gennaio"],
                "weeks": [],
                "solver_timeout": "abc",
            })

    @unittest.skipUnless(_TKINTER_OK, "tkinter non disponibile nel test environment")
    def test_perform_restart_rebuilds_step1_without_default_rows(self):
        app = TurniApp.__new__(TurniApp)
        build_calls = []
        shown_steps = []

        class _DummyLabel:
            def __init__(self):
                self.kw = None
            def config(self, **kw):
                self.kw = kw

        app.weeks_data = [{"week": "Sett.1"}]
        app._last_solver = object()
        app._last_result_text = "abc"
        app._pending_restart = True
        app._result_subtitle = _DummyLabel()
        app._set_result = lambda text: setattr(app, "_result_text_value", text)
        app.mesi = ["Gennaio"]
        app._build_step1_content = lambda add_default_row=True: build_calls.append(add_default_row)
        app._show_step = lambda step: shown_steps.append(step)

        TurniApp._perform_restart(app)

        self.assertEqual(app.weeks_data, [])
        self.assertIsNone(app._last_solver)
        self.assertEqual(app._last_result_text, "")
        self.assertFalse(app._pending_restart)
        self.assertEqual(build_calls, [False])
        self.assertEqual(shown_steps, [0])

    def test_target_file_lock_cleans_up_when_lock_write_fails(self):
        with tempfile.TemporaryDirectory() as tmp:
            target = pathlib.Path(tmp) / "sessione.json"
            real_write = os.write

            def flaky_write(fd, data):
                if data:
                    raise OSError("write failure")
                return real_write(fd, data)

            import turni.io_utils as io_utils_module
            with mock.patch.object(io_utils_module.os, "write", side_effect=flaky_write):
                with self.assertRaises(OSError):
                    with _TargetFileLock(str(target)):
                        pass

            self.assertFalse(pathlib.Path(str(target) + ".lock").exists())


if __name__ == "__main__":
    unittest.main()
