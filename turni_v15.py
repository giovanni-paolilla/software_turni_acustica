"""
Gestione Turni - Versione Definitiva v15
========================================
Fix applicati rispetto alla v14 (v15):
  [ROB]     Step 0 / sessione: almeno 2 operatori richiesti gia' in validazione
  [ROB]     Lock stale: PID vivo non viene piu' sottratto solo per anzianita';
              rilevamento pid reuse migliorato con token di start processo
  [ARCH]    Centralizzata anche la validazione di schedulabilita' minima
              (meno di 2 operatori disponibili) tra UI e sessione
  [TEST]    Aggiunti test su prerequisiti solver e lock con token processo

Fix applicati rispetto alla v13 (v14):
  [CRITICO] Nuovo calcolo: restart davvero pulito, senza settimane placeholder
              salvabili dallo step 0
  [ROB]     Lock file: cleanup sicuro se la creazione/scrittura del lock fallisce
  [ROB]     Sessione: solver_timeout validato in modo forte anche in load/save
  [ARCH]    Validazione settimane condivisa tra UI e sessione
  [TEST]    Aggiunti test su restart, timeout sessione e cleanup lock

Fix applicati rispetto alla v12 (v13):
  [CRITICO] Sessione: salvataggio round-trip safe e coerente con i dati
              realmente confermati/ricaricabili
  [ROB]     load_session: rifiuta sessioni con operatori o mesi vuoti
  [ROB]     Lock file: recupero automatico dei lock stale tramite PID/timestamp
  [ARCH]    Validazione sessione centralizzata in helper condivisi
  [TEST]    Aggiunti test su validazione sessione e stale lock

Fix applicati rispetto alla v11 (v12):
  [CRITICO] Annullamento: distinto correttamente da un completamento normale
              anche se esiste una soluzione parziale trovata prima dello stop
  [ROB]     load_session: rifiuta operatori/mesi/settimane vuoti o solo spazi
  [ROB]     load_session: rifiuta settimane duplicate nello stesso mese
  [ROB]     Scritture coordinate tra istanze tramite lock file per target
  [ARCH]    Regole di validazione centralizzate in helper condivisi
  [TEST]    Estesa la suite test per lock file e validazioni condivise

Fix applicati rispetto alla v10 (v11):
  [CRITICO] TurniSolver.solve(): annullamento onorato anche se richiesto prima
              dell'aggancio del CpSolver o subito prima della Solve()
  [ROB]     load_session: ripristino settimane per mese normalizzato coerente con
              la validazione (niente perdite silenziose per maiuscole/spazi)
  [ROB]     Step 0 / sessione JSON: mesi deduplicati o rifiutati se duplicati
              dopo normalizzazione (evita collisioni in _month_frames)
  [ARCH]    TurniSolver: reset esplicito dello stato interno a ogni solve()
  [TEST]    Estesa la suite test; discovery compatibile anche con
              "python -m unittest" dalla root

Fix applicati rispetto alla v9 (v10):
  [CRITICO] Solver: chiavi interne univoche per settimana (niente collisioni con
              settimane omonime nello stesso mese)
  [CRITICO] Annullamento: richiesta di stop inviata anche direttamente al CpSolver
              quando disponibile; close/restart ora attendono la fine del worker
  [SEC]     Export CSV sanificato contro formula injection
  [ROB]     load_session: valida radice JSON, operatori duplicati dopo
              normalizzazione, mesi incoerenti e busy_indices fuori range
  [ROB]     Salvataggi TXT/CSV/JSON resi atomici (tmp + replace)
  [UX]      "Ricomincia" rinominato in "Nuovo calcolo" e reso coerente con un
              reset reale dello step settimane
  [TEST]    Aggiunti test automatici per helper critici e requirements.txt

Fix applicati rispetto alla v8:
  [BUG]     _load_session: elementi di s["weeks"] validati come dict prima di
              issubset (lista di non-dict causava TypeError)
  [BUG]     _step0_next: aggiunta guardia su _solving (navigazione manuale tra
              step poteva ricostuire step1 mentre il thread era ancora vivo)
  [ARCH]    _month_frames inizializzato in __init__ (attributo di istanza mancante)
  [ARCH]    _load_session: rd = self._week_rows[-1] al posto di list comprehension
              O(n²) — la riga appena aggiunta è sempre l'ultima
  [ARCH]    solve(): ind da dict a list (chiavi intere consecutive → lista)
  [ARCH]    Type hints modernizzati: List/Dict/Optional → list/dict/X|None
              (stile Python 3.10+, con from __future__ import annotations)
  [ARCH]    _darken(): validazione lunghezza input con ValueError esplicito

Fix applicati rispetto alla v7 (v8):
  [BUG]     _run_solver: deep copy delle liste "available" e "busy" nello snapshot
              (shallow copy lasciava riferimenti condivisi col thread principale)
  [BUG]     _load_session: guard su _solving (caricamento durante calcolo = stato incoerente)
  [BUG]     _export e _save_session: try/except su I/O file (disco pieno, permessi)
  [BUG]     _step1_next: weeks_data costruito in variabile locale, assegnato solo a successo
  [ARCH]    ACCENT2 e SUCCESS2 rimossi (dead code dopo introduzione di _darken in v3)
  [ARCH]    _restart: resetta _last_solver e _last_result_text (evita export dati obsoleti)
  [ARCH]    _build_step1_content: parametro add_default_row elimina flusso
              costruisci-poi-distruggi in _load_session
  [ARCH]    _load_session: validazione tipi valori JSON (non solo chiavi)
  [UX]      Timeout salvato e ripristinato nella sessione JSON
  [UX]      _display_result: colore subtitle diverso in caso di annullamento

Fix applicati rispetto alla v3 (v4):
  [BUG]     Worker thread: try/except attorno a solve() — _solving non resta
              bloccato a True in caso di eccezione imprevista
  [BUG]     _export: rimosso dead code (if not self._last_solver) ora coperto
              da solver_ok
  [BUG]     _save_session: guard durante calcolo — impedisce salvataggio
              di stato inconsistente
  [ARCH]    Type hints su TurniSolver (metodi pubblici e __init__)
  [ARCH]    logging.getLogger per tracciabilita' senza print
  [UX]      SOLVER_TIMEOUT configurabile dall'utente in Step 0 (campo dedicato)
  [UX]      Rinomina n -> norm in _step0_next (nome ambiguo vs step number)

Fix applicati rispetto alla v2 (v3):
  [BUG]     normalize_name() in solve() per match corretto dei nomi busy
  [BUG]     Scrollbar _ops_text collegata al Text widget (yscrollcommand)
  [BUG]     Guardia doppio avvio solver in _step1_next (flag _solving)
  [BUG]     _display_result: guard winfo_exists() prima di aggiornare UI
  [BUG]     Anno default: anno corrente da datetime invece di hardcoded 2025
  [ARCH]    solved_ok su TurniSolver: rilevamento successo esplicito (non startswith)
  [ARCH]    _export usa solved_ok coerentemente con _display_result
  [ARCH]    _load_session: validazione struttura JSON (chiavi obbligatorie)
  [UX]      styled_btn: hover generico su tutti i colori (darken -15%)
  [UX]      ScrollableFrame._unbind_scroll: no unbind se mouse ancora nel canvas

Fix applicati rispetto alla v1 (v2):
  [CRITICO] Race condition: snapshot immutabili passati al thread solver
  [CRITICO] WM_DELETE_WINDOW: gestione chiusura durante calcolo
  [CRITICO] _last_result inizializzato in __init__
  [UX]      Navigazione Step1->Step0: preserva dati se operatori/mesi invariati
  [UX]      Pulsante Annulla durante il calcolo (via CpSolverSolutionCallback)
  [UX]      Rinumerazione settimane dopo eliminazione
  [UX]      MouseWheel vincolato al canvas (non bind_all globale)
  [ARCH]    TurniSolver separato dalla UI, testabile in isolamento
  [ARCH]    Pesi come costanti documentate a livello modulo
  [ARCH]    Salvataggio/caricamento sessione in JSON
  [EXTRA]   Validazione anno (intero 4 cifre)
  [EXTRA]   Export CSV oltre al TXT
  [EXTRA]   ttk.Style configurato per coerenza visiva
"""
from __future__ import annotations   # [v9 ARCH] type hints moderni su Python 3.8+

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import datetime
import json
import csv
import io
import re
import logging
import platform
import os
import tempfile
import time
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

try:
    from ortools.sat.python import cp_model
    ORTOOLS_OK = True
except ImportError:
    ORTOOLS_OK = False


# ══════════════════════════════════════════════════════════════
#  COSTANTI SOLVER
#  Peso equita': penalizza fortemente la differenza max-min turni.
#  100x > peso_penalty garantisce che l'equita' abbia sempre precedenza
#  sulla preferenza di non riusare l'operatore sab+mart nella stessa sett.
# ══════════════════════════════════════════════════════════════
PESO_DIFF        = 100
PESO_PENALTY     = 1
SOLVER_TIMEOUT   = 60.0
LOCK_STALE_SECONDS = 24 * 60 * 60


# ══════════════════════════════════════════════════════════════
#  FONT  [v7 ARCH] cross-platform: Segoe UI non esiste su macOS/Linux
# ══════════════════════════════════════════════════════════════
_SYS = platform.system()
_FONT_FAMILY = ("Segoe UI"    if _SYS == "Windows" else
                "SF Pro Text" if _SYS == "Darwin"  else
                "DejaVu Sans")


# ══════════════════════════════════════════════════════════════
#  TEMA
# ══════════════════════════════════════════════════════════════
BG           = "#F0F4F8"
PANEL_BG     = "#FFFFFF"
ACCENT       = "#2563EB"
# ACCENT2 / SUCCESS2 rimossi in v5: dead code dopo _darken() introdotto in v3
SUCCESS      = "#16A34A"
DANGER       = "#DC2626"
WARN         = "#D97706"
MUTED        = "#64748B"
TEXT         = "#1E293B"
BORDER       = "#CBD5E1"
HEADER_BG    = "#1E3A5F"
# [v6 ARCH] colori inline promossi a costanti di tema
MONTH_HDR_BG = "#EFF6FF"   # sfondo header sezione mese nello step 1
ACCENT_LIGHT = "#DBEAFE"   # sfondo pulsante "+ Aggiungi" e checkbox selezionato
BTN_SECONDARY= "#E2E8F0"   # pulsanti navigazione secondari (Indietro, Modifica)
DANGER_LIGHT = "#FEE2E2"   # sfondo pulsante "X" elimina settimana
HEADER_BTN   = "#1E4A7A"   # pulsanti Salva/Carica sessione nell'header
INFO_WARN    = "#FCA5A5"   # testo avviso ortools mancante nell'header
CYAN         = "#0891B2"   # pulsante "Salva CSV"

FONT_TITLE = (_FONT_FAMILY, 18, "bold")
FONT_LABEL = (_FONT_FAMILY, 10)
FONT_BOLD  = (_FONT_FAMILY, 10, "bold")
FONT_SMALL = (_FONT_FAMILY, 9)
FONT_MONO  = ("Courier New", 10)


def configure_ttk_style() -> None:    # [v7 ARCH] type hint
    s = ttk.Style()
    s.theme_use("clam")
    s.configure("Vertical.TScrollbar",
                background=BORDER, troughcolor=BG,
                arrowcolor=MUTED, bordercolor=BG)
    s.configure("Horizontal.TScrollbar",
                background=BORDER, troughcolor=BG,
                arrowcolor=MUTED, bordercolor=BG)
    s.configure("TProgressbar",
                troughcolor=BG, background=ACCENT,
                darkcolor=ACCENT, lightcolor=ACCENT)


# ══════════════════════════════════════════════════════════════
#  HELPER UI
# ══════════════════════════════════════════════════════════════
def normalize_name(name: str) -> str:           # [v6 ARCH] type hints funzioni modulo
    return " ".join(name.strip().split()).lower()


def _safe_csv_cell(value: str) -> str:
    """Neutralizza formule potenzialmente eseguibili nei fogli di calcolo."""
    if not isinstance(value, str):
        return value
    stripped = value.lstrip()
    if stripped and stripped[0] in ("=", "+", "-", "@"):
        return "'" + value
    return value


def _dedupe_normalized_texts(values: list[str]) -> tuple[list[str], list[str], list[str]]:
    """Deduplica stringhe preservando l'ordine e confrontando via normalize_name()."""
    unique_values: list[str] = []
    unique_norms: list[str] = []
    duplicates: list[str] = []
    for value in values:
        norm = normalize_name(value)
        if norm not in unique_norms:
            unique_values.append(value)
            unique_norms.append(norm)
        else:
            duplicates.append(value)
    return unique_values, unique_norms, duplicates


def _group_weeks_by_normalized_month(weeks: list[dict]) -> dict[str, list[dict]]:
    """Raggruppa le settimane usando una chiave mese normalizzata."""
    grouped: dict[str, list[dict]] = {}
    for week in weeks:
        grouped.setdefault(normalize_name(week["month"]), []).append(week)
    return grouped


def _order_weeks_by_declared_months(weeks: list[dict], declared_months: list[str]) -> list[dict]:
    """Ordina le settimane per mese dichiarato, preservando l'ordine interno al mese.

    Evita output interleavati del tipo Gennaio1, Febbraio1, Marzo1, Gennaio2...
    quando la UI o il caricamento sessione hanno costruito la lista in quell'ordine.
    """
    grouped = _group_weeks_by_normalized_month(weeks)
    ordered: list[dict] = []
    seen_months: set[str] = set()

    for month in declared_months:
        month_norm = normalize_name(month)
        if month_norm in grouped:
            ordered.extend(grouped[month_norm])
            seen_months.add(month_norm)

    for week in weeks:
        month_norm = normalize_name(week["month"])
        if month_norm not in seen_months:
            ordered.append(week)
            seen_months.add(month_norm)

    return ordered


def _find_blank_text_entries(values: list[str]) -> list[str]:
    """Restituisce gli elementi vuoti o composti solo da spazi."""
    return [value for value in values if not normalize_name(value)]


def _normalized_week_key(month: str, week: str) -> tuple[str, str]:
    """Chiave canonica di una settimana: mese + nome settimana normalizzati."""
    return normalize_name(month), normalize_name(week)


def _find_duplicate_week_keys(weeks: list[dict]) -> list[tuple[str, str]]:
    """Individua settimane duplicate nello stesso mese dopo normalizzazione."""
    duplicates: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()
    dup_seen: set[tuple[str, str]] = set()
    for week in weeks:
        key = _normalized_week_key(week["month"], week["week"])
        if key in seen and key not in dup_seen:
            duplicates.append((week["month"], week["week"]))
            dup_seen.add(key)
        else:
            seen.add(key)
    return duplicates


class SessionValidationError(ValueError):
    """Errore di validazione dei dati di sessione."""


def _parse_operatori_text(raw: str) -> list[str]:
    return [x.strip() for part in raw.splitlines() for x in part.split(",") if x.strip()]


def _parse_mesi_text(raw: str) -> list[str]:
    return [m.strip() for m in raw.split(",") if m.strip()]


def _canonicalize_timeout_text(timeout_value, *, error_message: str) -> str:
    timeout_text = str(timeout_value).strip()
    try:
        timeout_val = float(timeout_text)
        if timeout_val <= 0:
            raise ValueError
    except (TypeError, ValueError) as exc:
        raise SessionValidationError(error_message) from exc
    return str(timeout_val)


def _validate_operator_pool(operatori: list[str], *, error_message: str) -> None:
    if len(operatori) < 2:
        raise SessionValidationError(error_message)


def _week_available_count(week: dict, operator_count: int) -> int:
    if "available" in week:
        available = week.get("available")
        if isinstance(available, list):
            return len(set(available))
    busy_indices = week.get("busy_indices")
    if isinstance(busy_indices, list):
        return operator_count - len(set(busy_indices))
    busy_names = week.get("busy")
    if isinstance(busy_names, list):
        return operator_count - len(set(busy_names))
    return operator_count


def _validate_solver_ready_weeks(weeks: list[dict], *, operator_count: int,
                                 error_prefix: str = "File sessione non valido") -> None:
    if operator_count < 2:
        raise SessionValidationError(
            f"{error_prefix}: servono almeno 2 operatori totali per pianificare i turni."
        )
    for week in weeks:
        if _week_available_count(week, operator_count) < 2:
            raise SessionValidationError(
                f"{error_prefix}: '{week['week']}' ({week['month']}): meno di 2 operatori disponibili."
            )


def _validate_week_entries(weeks: list[dict], *, declared_months: list[str] | None = None,
                           error_prefix: str = "File sessione non valido") -> list[dict]:
    """Valida e canonizza le settimane condividendo le regole tra UI e sessione."""
    declared_months_norm = None
    if declared_months is not None:
        declared_months_norm = {normalize_name(m) for m in declared_months}

    canonical_weeks: list[dict] = []
    for week in weeks:
        month = week.get("month")
        week_name = week.get("week")

        if not isinstance(month, str) or not normalize_name(month):
            raise SessionValidationError(
                f"{error_prefix}: presenti settimane associate a mesi non dichiarati nella sessione."
            )
        if declared_months_norm is not None and normalize_name(month) not in declared_months_norm:
            raise SessionValidationError(
                f"{error_prefix}: presenti settimane associate a mesi non dichiarati nella sessione."
            )
        if not isinstance(week_name, str):
            raise SessionValidationError(
                f"{error_prefix}: il nome di ogni settimana deve essere testo."
            )
        if not normalize_name(week_name):
            raise SessionValidationError(
                f"{error_prefix}: presenti settimane con nome vuoto o solo spazi."
            )

        canonical_weeks.append({
            **week,
            "month": month.strip(),
            "week": week_name.strip(),
        })

    duplicate_weeks = _find_duplicate_week_keys(canonical_weeks)
    if duplicate_weeks:
        dup_month, dup_week = duplicate_weeks[0]
        raise SessionValidationError(
            f"{error_prefix}: settimana duplicata '{dup_week}' nel mese '{dup_month}'."
        )

    return canonical_weeks


def _parse_step0_inputs(anno: str, mesi_text: str, operatori_text: str, timeout_text: str
                        ) -> tuple[str, str, list[str], list[str], list[str], list[str], list[str]]:
    """Parsa e valida i campi dello step 0 restituendo una forma canonica."""
    anno = anno.strip()
    if not re.fullmatch(r"\d{4}", anno):
        raise SessionValidationError("Anno non valido. Inserisci 4 cifre (es. 2025).")

    timeout_text = _canonicalize_timeout_text(
        timeout_text,
        error_message="Timeout non valido. Inserisci un numero positivo (es. 60).",
    )

    raw_ops = _parse_operatori_text(operatori_text)
    if not raw_ops:
        raise SessionValidationError("Inserisci almeno un operatore.")
    if _find_blank_text_entries(raw_ops):
        raise SessionValidationError(
            "Gli operatori non possono essere vuoti o composti solo da spazi."
        )

    raw_mesi = _parse_mesi_text(mesi_text)
    if not raw_mesi:
        raise SessionValidationError("Inserisci almeno un mese.")
    if _find_blank_text_entries(raw_mesi):
        raise SessionValidationError(
            "I mesi non possono essere vuoti o composti solo da spazi."
        )

    new_ops, new_norm, dup_ops = _dedupe_normalized_texts(raw_ops)
    new_mesi, _, dup_mesi = _dedupe_normalized_texts(raw_mesi)
    _validate_operator_pool(
        new_ops,
        error_message=(
            "Servono almeno 2 operatori totali per pianificare i turni "
            "(giovedi Audio e Video devono essere assegnati a persone diverse)."
        ),
    )
    return anno, timeout_text, new_ops, new_norm, new_mesi, dup_ops, dup_mesi


def _validate_session_payload(session: dict) -> dict:
    """Valida e canonizza il payload di sessione usato sia in save che in load."""
    if not isinstance(session, dict):
        raise SessionValidationError(
            "File sessione non valido: la radice JSON deve essere un oggetto."
        )

    required_session_keys = {"anno", "operatori", "mesi", "weeks"}
    if not required_session_keys.issubset(session):
        missing = required_session_keys - session.keys()
        raise SessionValidationError(
            f"File sessione non valido.\nCampi mancanti: {', '.join(sorted(missing))}"
        )

    ops = session["operatori"]
    mesi = session["mesi"]
    weeks = session["weeks"]

    if (not isinstance(ops, list) or not all(isinstance(o, str) for o in ops)
            or not isinstance(mesi, list)
            or not all(isinstance(m, str) for m in mesi)):
        raise SessionValidationError(
            "File sessione non valido: 'operatori' e 'mesi' devono essere liste di testo."
        )

    ops = [o.strip() for o in ops]
    mesi = [m.strip() for m in mesi]

    if not ops:
        raise SessionValidationError("File sessione non valido: inserisci almeno un operatore.")
    if not mesi:
        raise SessionValidationError("File sessione non valido: inserisci almeno un mese.")

    if _find_blank_text_entries(ops):
        raise SessionValidationError(
            "File sessione non valido: operatori vuoti o composti solo da spazi."
        )
    if _find_blank_text_entries(mesi):
        raise SessionValidationError(
            "File sessione non valido: mesi vuoti o composti solo da spazi."
        )

    _, _, dup_ops = _dedupe_normalized_texts(ops)
    if dup_ops:
        raise SessionValidationError(
            "File sessione non valido: operatori duplicati dopo normalizzazione.\n"
            + "\n".join(dup_ops)
        )

    _validate_operator_pool(
        ops,
        error_message=(
            "File sessione non valido: servono almeno 2 operatori totali per "
            "pianificare i turni."
        ),
    )

    _, _, dup_mesi = _dedupe_normalized_texts(mesi)
    if dup_mesi:
        raise SessionValidationError(
            "File sessione non valido: mesi duplicati dopo normalizzazione.\n"
            + "\n".join(dup_mesi)
        )

    if not isinstance(weeks, list):
        raise SessionValidationError("File sessione non valido: 'weeks' deve essere una lista.")
    if not all(isinstance(w, dict) for w in weeks):
        raise SessionValidationError(
            "File sessione non valido: ogni settimana deve essere un oggetto."
        )

    required_week_keys = {"month", "week", "busy_indices"}
    bad_weeks = [w for w in weeks if not required_week_keys.issubset(w)]
    if bad_weeks:
        raise SessionValidationError(
            f"File sessione non valido: {len(bad_weeks)} settimana/e malformata/e."
        )

    canonical_weeks = _validate_week_entries(
        weeks,
        declared_months=mesi,
        error_prefix="File sessione non valido",
    )
    for week in canonical_weeks:
        busy_indices = week.get("busy_indices", [])
        if (not isinstance(busy_indices, list)
                or not all(isinstance(i, int) and i >= 0 for i in busy_indices)):
            raise SessionValidationError(
                f"busy_indices non valido nella settimana '{week['week']}'.\nDevono essere interi non negativi."
            )
        week["busy_indices"] = list(busy_indices)

    too_large_idx = [
        (w["week"], idx)
        for w in canonical_weeks
        for idx in w["busy_indices"]
        if idx >= len(ops)
    ]
    if too_large_idx:
        week_name, idx = too_large_idx[0]
        raise SessionValidationError(
            f"busy_indices fuori range nella settimana '{week_name}' (indice {idx})."
        )

    _validate_solver_ready_weeks(
        canonical_weeks,
        operator_count=len(ops),
        error_prefix="File sessione non valido",
    )

    anno_val = str(session.get("anno", "")).strip()
    if not re.fullmatch(r"\d{4}", anno_val):
        raise SessionValidationError(
            f"Anno nella sessione non valido: '{anno_val}' (atteso formato YYYY)."
        )

    validated = {
        "anno": anno_val,
        "operatori": ops,
        "mesi": mesi,
        "weeks": canonical_weeks,
    }
    if "solver_timeout" in session:
        validated["solver_timeout"] = _canonicalize_timeout_text(
            session["solver_timeout"],
            error_message="File sessione non valido: solver_timeout deve essere un numero positivo.",
        )
    return validated


def _pid_is_alive(pid: int) -> bool:
    if pid <= 0:
        return False
    try:
        os.kill(pid, 0)
    except ProcessLookupError:
        return False
    except PermissionError:
        return True
    except OSError:
        return False
    else:
        return True


def _read_lock_metadata(lock_path: str) -> dict | None:
    try:
        with open(lock_path, encoding="utf-8") as f:
            payload = json.load(f)
    except (OSError, json.JSONDecodeError, TypeError, ValueError):
        return None
    return payload if isinstance(payload, dict) else None


def _get_process_start_token(pid: int) -> str | None:
    if pid <= 0:
        return None
    proc_stat = f"/proc/{pid}/stat"
    try:
        with open(proc_stat, encoding="utf-8") as f:
            stat = f.read()
    except OSError:
        return None

    try:
        after_comm = stat.rsplit(")", 1)[1].strip()
        fields = after_comm.split()
        return fields[19]
    except (IndexError, ValueError):
        return None


def _is_stale_lock(lock_path: str, *, stale_after_seconds: int = LOCK_STALE_SECONDS) -> bool:
    metadata = _read_lock_metadata(lock_path)
    now = time.time()

    if metadata is not None:
        pid = metadata.get("pid")
        created_at = metadata.get("created_at")
        proc_start_token = metadata.get("proc_start_token")

        if isinstance(pid, int):
            if not _pid_is_alive(pid):
                return True

            current_token = _get_process_start_token(pid)
            if (proc_start_token is not None and current_token is not None
                    and str(proc_start_token) != str(current_token)):
                return True

            return False

        if isinstance(created_at, (int, float)) and now - float(created_at) > stale_after_seconds:
            return True
        return False

    try:
        mtime = os.path.getmtime(lock_path)
    except OSError:
        return False
    return now - mtime > stale_after_seconds


class _TargetFileLock:
    """Lock file per coordinare scritture concorrenti sullo stesso target."""

    def __init__(self, target_path: str) -> None:
        self.target_path = os.path.abspath(target_path)
        self.lock_path = self.target_path + ".lock"
        self._fd: int | None = None

    def _try_acquire(self) -> bool:
        try:
            fd = os.open(self.lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
        except FileExistsError:
            return False

        try:
            payload = json.dumps({
                "pid": os.getpid(),
                "created_at": time.time(),
                "proc_start_token": _get_process_start_token(os.getpid()),
            })
            os.write(fd, payload.encode("utf-8", errors="ignore"))
            os.fsync(fd)
        except Exception:
            try:
                os.close(fd)
            except OSError:
                pass
            try:
                os.unlink(self.lock_path)
            except OSError:
                pass
            raise

        self._fd = fd
        return True

    def __enter__(self) -> _TargetFileLock:
        if self._try_acquire():
            return self

        if _is_stale_lock(self.lock_path):
            try:
                os.unlink(self.lock_path)
            except FileNotFoundError:
                pass
            except OSError:
                pass
            if self._try_acquire():
                return self

        raise BlockingIOError(
            f"File occupato da un'altra istanza: '{self.target_path}'."
        )

    def __exit__(self, exc_type, exc, tb) -> None:
        if self._fd is not None:
            try:
                os.close(self._fd)
            finally:
                self._fd = None
        try:
            os.unlink(self.lock_path)
        except FileNotFoundError:
            pass


def _write_text_file_atomic(path: str, data: str, newline: str | None = None) -> None:
    """Scrive testo in modo atomico tramite file temporaneo + os.replace()."""
    directory = os.path.dirname(os.path.abspath(path)) or "."
    with _TargetFileLock(path):
        fd, tmp_path = tempfile.mkstemp(prefix=".tmp_turni_", dir=directory)
        try:
            with os.fdopen(fd, "w", encoding="utf-8", newline=newline) as f:
                f.write(data)
                f.flush()
                os.fsync(f.fileno())
            os.replace(tmp_path, path)
        except Exception:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
            raise




def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False,
                   align: int = WD_ALIGN_PARAGRAPH.LEFT,
                   font_size: int = 10, font_name: str = "Aptos",
                   color: RGBColor | None = None) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    if color is not None:
        font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_row_min_height(row, height_cm: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    twips = str(int((height_cm / 2.54) * 1440))
    tr_height.set(qn("w:val"), twips)
    tr_height.set(qn("w:hRule"), "atLeast")


def _format_docx_table(table, column_widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(column_widths_cm):
            row.cells[idx].width = Cm(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "110")
                node.set(qn("w:type"), "dxa")


def _extract_week_dates_labels(week_label: str) -> tuple[str, str]:
    """Estrae le etichette DATA in formato breve usato nel DOCX finale.

    Esempi supportati:
      - "Sett. 08-11 Gen"          -> ("08 Gen", "11 Gen")
      - "Sett. 30 Apr - 03 Mag"    -> ("30 Apr", "03 Mag")

    In fallback usa i soli numeri trovati; se il parsing fallisce evita crash e
    restituisce etichette generiche mantenendo il documento esportabile.
    """
    label = (week_label or "").strip()
    compact = re.sub(r"\s*-\s*", "-", label)
    match = re.search(
        r"(?i)(\d{1,2})\s*([A-Za-zÀ-ÿ]{3,})?\s*-\s*(\d{1,2})\s*([A-Za-zÀ-ÿ]{3,})?",
        compact,
    )
    if match:
        d1, m1, d2, m2 = match.groups()
        if not m1 and m2:
            m1 = m2
        if not m2 and m1:
            m2 = m1
        left = d1.zfill(2) + (f" {m1[:3].title()}" if m1 else "")
        right = d2.zfill(2) + (f" {m2[:3].title()}" if m2 else "")
        return left, right

    nums = re.findall(r"\b(\d{1,2})\b", label)
    if len(nums) >= 2:
        return nums[0].zfill(2), nums[1].zfill(2)
    if len(nums) == 1:
        return nums[0].zfill(2), ""
    return (label or "—", "—")


def build_turni_docx(path: str, anno: str, result_rows: list[dict], operatori: list[str],
                     counts: list[int], status_str: str, diff_val: int, penalty_val: int) -> None:
    """Genera il DOCX con il layout grafico continuo richiesto dall'utente."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(0.75)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.55)
    section.right_margin = Cm(0.55)

    page_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm

    base_style = doc.styles["Normal"]
    base_style.font.name = "Arial"
    base_style.font.size = Pt(10)

    ordered_months = list(dict.fromkeys(row["month"] for row in result_rows))
    months_caption = " - ".join(ordered_months) if ordered_months else "Programmazione"

    title_table = doc.add_table(rows=1, cols=1)
    title_table.style = "Table Grid"
    title_table.autofit = False
    title_cell = title_table.rows[0].cells[0]
    title_cell.width = Cm(page_width_cm)
    _set_cell_shading(title_cell, "4C79C5")
    _clear_cell(title_cell)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = title_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(1)
    r1 = p1.add_run("AUDIO/VIDEO MESSINA-GANZIRRI")
    r1.bold = True
    r1.italic = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = title_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"{months_caption}   {anno}")
    r2.bold = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0xA8, 0xD0, 0x33)
    _set_row_min_height(title_table.rows[0], 1.45)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _format_docx_table(table, [4.65, 7.0, 7.0])

    hdr = table.rows[0].cells
    headers = ["DATA", "VIDEO", "AUDIO"]
    for idx, value in enumerate(headers):
        _set_cell_text(hdr[idx], value, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14,
                       font_name="Times New Roman",
                       color=RGBColor(0x2D, 0x3E, 0x56))
        _set_cell_shading(hdr[idx], "C7D5EC")
    _set_row_min_height(table.rows[0], 0.95)

    body_fill = "E9E9E9"
    for row_data in result_rows:
        left_label, right_label = _extract_week_dates_labels(row_data.get("week", ""))

        giovedi = table.add_row().cells
        giovedi_values = [
            left_label or "—",
            row_data.get("video", "") or "",
            row_data.get("audio", "") or "",
        ]
        for idx, value in enumerate(giovedi_values):
            _set_cell_text(
                giovedi[idx],
                str(value),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                font_size=10,
                font_name="Arial",
            )
            _set_cell_shading(giovedi[idx], body_fill)
        _set_row_min_height(table.rows[-1], 0.55)

        domenica = table.add_row().cells
        merged = domenica[1].merge(domenica[2])
        _set_cell_text(domenica[0], right_label or "—",
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       font_size=10, font_name="Arial")
        _set_cell_shading(domenica[0], body_fill)
        _set_cell_text(merged, row_data.get("sabato", "") or "",
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       font_size=10, font_name="Arial")
        _set_cell_shading(merged, body_fill)
        _set_row_min_height(table.rows[-1], 0.55)

    footer = section.footer.paragraphs[0]
    footer.clear()

    doc.save(path)


def _darken(hex_color: str, factor: float = 0.85) -> str:
    """Scurisce un colore esadecimale del (1-factor)*100 %.

    Raises:
        ValueError: se hex_color non è nel formato #RRGGBB.
    """
    h = hex_color.lstrip("#")
    if len(h) != 6:                       # [v9 ARCH] validazione lunghezza input
        raise ValueError(f"_darken: colore non valido '{hex_color}' (atteso #RRGGBB)")
    r, g, b = (int(h[i:i+2], 16) for i in (0, 2, 4))
    r, g, b = (max(0, int(c * factor)) for c in (r, g, b))
    return f"#{r:02x}{g:02x}{b:02x}"


def styled_btn(parent: tk.Widget, text: str, command,
               bg: str = ACCENT, fg: str = "white", **kw) -> tk.Button:
    # [v3 UX] hover generico: darken -15% su qualsiasi colore
    hover = _darken(bg)
    btn = tk.Button(parent, text=text, command=command,
                    bg=bg, fg=fg, font=FONT_BOLD,
                    relief="flat", cursor="hand2",
                    padx=14, pady=6, **kw)
    btn.bind("<Enter>", lambda e: btn.config(bg=hover))
    btn.bind("<Leave>", lambda e: btn.config(bg=bg))
    return btn


def card(parent: tk.Widget, **kw) -> tk.Frame:
    return tk.Frame(parent, bg=PANEL_BG, relief="flat",
                    highlightbackground=BORDER, highlightthickness=1, **kw)


# ══════════════════════════════════════════════════════════════
#  SCROLLABLE FRAME  [UX fix: niente bind_all globale]
# ══════════════════════════════════════════════════════════════
class ScrollableFrame(tk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, bg=BG, **kw)
        self._canvas = tk.Canvas(self, bg=BG, highlightthickness=0)
        sb = ttk.Scrollbar(self, orient="vertical",
                           command=self._canvas.yview)
        self.inner = tk.Frame(self._canvas, bg=BG)
        self.inner.bind("<Configure>",
            lambda e: self._canvas.configure(
                scrollregion=self._canvas.bbox("all")))
        self._canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self._canvas.configure(yscrollcommand=sb.set)
        self._canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        # Scroll solo quando il mouse e' dentro il frame
        self.inner.bind("<Enter>", self._bind_scroll)
        self.inner.bind("<Leave>", self._unbind_scroll)

    def _on_scroll(self, e):
        # [v6 BUG] fix cross-platform:
        #   Linux:   usa e.num (Button-4 / Button-5), e.delta è sempre 0
        #   macOS:   e.delta è float (es. 3.5), //120 dava 0 → scroll non funzionava
        #   Windows: e.delta è multiplo di 120, solo il segno conta
        # [v8 BUG] e.delta==0 su Windows (trackpad in stato intermedio) → no-op
        if e.num == 4:
            delta = -1
        elif e.num == 5:
            delta = 1
        elif e.delta > 0:
            delta = -1
        elif e.delta < 0:
            delta = 1
        else:
            return   # e.delta == 0: nessuno scroll
        self._canvas.yview_scroll(delta, "units")

    def _bind_scroll(self, _=None):
        self._canvas.bind_all("<MouseWheel>", self._on_scroll)
        self._canvas.bind_all("<Button-4>",   self._on_scroll)
        self._canvas.bind_all("<Button-5>",   self._on_scroll)

    def _unbind_scroll(self, event=None):
        # [v3 UX] Evita unbind se il cursore e' ancora sopra il canvas
        # (accade quando si entra in un widget figlio di self.inner)
        if event is not None:
            w = self._canvas.winfo_containing(event.x_root, event.y_root)
            if w and (w == self._canvas or str(w).startswith(str(self._canvas))):
                return
        self._canvas.unbind_all("<MouseWheel>")
        self._canvas.unbind_all("<Button-4>")
        self._canvas.unbind_all("<Button-5>")


# ══════════════════════════════════════════════════════════════
#  SOLVER  [ARCH: separato dalla UI, testabile in isolamento]
# ══════════════════════════════════════════════════════════════
class _CancelCallback(cp_model.CpSolverSolutionCallback if ORTOOLS_OK else object):
    """Ferma la ricerca quando cancel_event e' impostato."""
    def __init__(self, cancel_event):
        if ORTOOLS_OK:
            super().__init__()
        self._ev = cancel_event

    def on_solution_callback(self):
        if self._ev.is_set():
            self.StopSearch()


class TurniSolver:
    """Logica CP-SAT completamente disaccoppiata dalla UI.

    Attributi pubblici dopo solve():
        result_rows  -- lista di dict con month/week/audio/video/sabato/busy
        counts       -- turni totali per operatore (indice parallelo a operatori)
        diff_val     -- delta max-min (obiettivo di equita')
        penalty_val  -- penalita' sabato sovrapposto a turno settimanale
        status_str   -- "OTTIMALE" | "FATTIBILE (timeout)"
        solved_ok    -- True solo se solve() ha trovato una soluzione valida
    """

    def __init__(self, operatori: list[str], weeks_data: list[dict]) -> None:
        self.operatori:   list[str]  = operatori
        self.weeks_data:  list[dict] = weeks_data
        self.result_rows: list[dict] = []
        self.counts:      list[int]  = []
        self.diff_val:    int        = 0
        self.penalty_val: int        = 0
        self.status_str:  str        = ""
        self.solved_ok:   bool       = False   # [v3 ARCH] flag esplicito di successo
        self.cancelled:   bool       = False
        self.partial_result_available: bool = False
        self.cp_solver                = None

    def _reset_result_state(self) -> None:
        self.result_rows = []
        self.counts = []
        self.diff_val = 0
        self.penalty_val = 0
        self.status_str = ""
        self.solved_ok = False
        self.cancelled = False
        self.partial_result_available = False
        self.cp_solver = None

    def solve(self, cancel_event: threading.Event | None = None,
              timeout: float = SOLVER_TIMEOUT) -> str:
        """Esegue il solver CP-SAT e restituisce il testo risultato.

        Args:
            cancel_event: Event impostato per interrompere la ricerca.
            timeout:      Limite di tempo in secondi (default: SOLVER_TIMEOUT).

        Returns:
            Stringa formattata del risultato, o messaggio di errore.
        """
        self._reset_result_state()

        if not ORTOOLS_OK:
            return ("ERRORE: ortools non installato.\n"
                    "Esegui:  pip install ortools")

        if cancel_event is None:
            cancel_event = threading.Event()
        if cancel_event.is_set():
            self.cancelled = True
            return "Calcolo annullato prima di trovare una soluzione."

        operatori  = self.operatori
        weeks_data = self.weeks_data
        n          = len(operatori)
        model      = cp_model.CpModel()
        turn_vars  = {}

        for week_idx, week in enumerate(weeks_data):
            key   = week_idx
            avail = week["available"]
            if not avail:
                return (f"Errore: nessun operatore per "
                        f"'{week['week']}' ({week['month']}).")
            ta  = model.NewIntVarFromDomain(
                cp_model.Domain.FromValues(avail), f"ta_w{key}")
            tv  = model.NewIntVarFromDomain(
                cp_model.Domain.FromValues(avail), f"tv_w{key}")
            sat = model.NewIntVarFromDomain(
                cp_model.Domain.FromValues(avail), f"sat_w{key}")
            model.Add(ta != tv)
            turn_vars[key] = {"t_audio": ta, "t_video": tv, "sat": sat}

        penalties = []
        for week_idx, week in enumerate(weeks_data):
            key = week_idx
            ta, tv, sat = (turn_vars[key]["t_audio"],
                           turn_vars[key]["t_video"],
                           turn_vars[key]["sat"])
            for t_var, lbl in [(ta, "a"), (tv, "v")]:
                pen = model.NewBoolVar(f"pen_{lbl}_w{key}")
                model.Add(sat == t_var).OnlyEnforceIf(pen)
                model.Add(sat != t_var).OnlyEnforceIf(pen.Not())
                penalties.append(pen)
        tot_pen = model.NewIntVar(0, 2*len(weeks_data), "tp")
        model.Add(tot_pen == sum(penalties))

        # [v9 ARCH] lista al posto di dict con chiavi intere consecutive (più efficiente)
        ind: list[list] = [[] for _ in range(n)]
        for week_idx, week in enumerate(weeks_data):
            key = week_idx
            for ruolo, var in turn_vars[key].items():
                for i in week["available"]:
                    b = model.NewBoolVar(f"b_w{key}_{ruolo}_{i}")
                    model.Add(var == i).OnlyEnforceIf(b)
                    model.Add(var != i).OnlyEnforceIf(b.Not())
                    ind[i].append(b)

        cnt_vars = []
        for i in range(n):
            c = model.NewIntVar(0, 3*len(weeks_data), f"cnt{i}")
            model.Add(c == sum(ind[i]))
            cnt_vars.append(c)

        mx = model.NewIntVar(0, 3*len(weeks_data), "mx")
        mn = model.NewIntVar(0, 3*len(weeks_data), "mn")
        for c in cnt_vars:
            model.Add(c <= mx)
            model.Add(c >= mn)
        diff = model.NewIntVar(0, 3*len(weeks_data), "diff")
        model.Add(diff == mx - mn)

        model.Minimize(diff * PESO_DIFF + tot_pen * PESO_PENALTY)

        solver = cp_model.CpSolver()
        solver.parameters.max_time_in_seconds = timeout  # [v4 UX] timeout configurabile
        self.cp_solver = solver
        if cancel_event.is_set():
            self.cancelled = True
            self.cp_solver = None
            return "Calcolo annullato prima di trovare una soluzione."
        try:
            status = solver.Solve(model, _CancelCallback(cancel_event))
            cancel_requested = cancel_event.is_set()
        finally:
            self.cp_solver = None

        if cancel_requested and status not in [
                cp_model.OPTIMAL, cp_model.FEASIBLE]:
            self.cancelled = True
            return "Calcolo annullato prima di trovare una soluzione."

        if status not in [cp_model.OPTIMAL, cp_model.FEASIBLE]:
            return "Nessuna soluzione trovata. Controlla i dati inseriti."

        self.result_rows = []
        for week_idx, week in enumerate(weeks_data):
            key = week_idx
            ta  = solver.Value(turn_vars[key]["t_audio"])
            tv  = solver.Value(turn_vars[key]["t_video"])
            sat = solver.Value(turn_vars[key]["sat"])
            # [v3 BUG] normalize_name invece di op.lower(): gestisce spazi multipli
            busy_names = ", ".join(
                operatori[i] for i, op in enumerate(operatori)
                if normalize_name(op) in week["busy"]
            ) if week["busy"] else ""
            self.result_rows.append({
                "month":  week["month"],
                "week":   week["week"],
                "audio":  operatori[ta],
                "video":  operatori[tv],
                "sabato": operatori[sat],
                "busy":   busy_names,
            })

        self.counts      = [solver.Value(cnt_vars[i]) for i in range(n)]
        self.diff_val    = solver.Value(diff)
        self.penalty_val = solver.Value(tot_pen)
        if cancel_requested:
            self.cancelled = True
            self.partial_result_available = True
            self.status_str = "ANNULLATO DALL'UTENTE (soluzione parziale)"
            return self._format_text()
        self.status_str  = ("OTTIMALE" if status == cp_model.OPTIMAL
                            else "FATTIBILE (timeout)")
        self.solved_ok   = True   # [v3 ARCH] segnale esplicito di successo
        return self._format_text()

    def _format_text(self) -> str:      # [v7 ARCH] type hint
        lines = []
        sep   = "-" * 64
        lines += [sep, "  PROGRAMMA TURNI", sep]

        # [v7 ARCH] larghezza colonna Audio calcolata dinamicamente (no hardcoded 20)
        audio_col_w: int = max((len(r["audio"]) for r in self.result_rows), default=20)

        cur_month = None
        for r in self.result_rows:
            if r["month"] != cur_month:
                cur_month = r["month"]
                lines.append(f"\n  {cur_month.upper()}")
                lines.append("-" * 40)
            lines.append(f"  {r['week']}")
            lines.append(
                f"    Giovedi  | Audio: {r['audio']:<{audio_col_w}}  Video: {r['video']}")
            lines.append(f"    Domenica | Operatore A/V: {r['sabato']}")
            if r["busy"]:
                lines.append(f"    Busy     | {r['busy']}")
            lines.append("")

        lines += [sep, "  RIEPILOGO TURNI PER OPERATORE", sep]
        paired  = sorted(zip(self.operatori, self.counts), key=lambda x: -x[1])
        # [v7 BUG] max() con default=10: evita ValueError se paired e' vuoto
        max_n   = max((len(p[0]) for p in paired), default=10)
        for op, cnt in paired:
            lines.append(f"  {op:<{max_n}}  {'|' * cnt}  ({cnt})")

        lines += [
            "",
            f"  Stato           : {self.status_str}",
            f"  Delta max-min   : {self.diff_val}",
            f"  Penalita' domenica: {self.penalty_val}",
            sep,
        ]
        return "\n".join(lines)

    def format_csv(self, anno: str) -> str:   # [v7 ARCH] type hint
        buf = io.StringIO()
        w   = csv.writer(buf)
        w.writerow(["Anno","Mese","Settimana",
                    "Giovedi Audio","Giovedi Video","Domenica A/V","Busy"])
        for r in self.result_rows:
            w.writerow([_safe_csv_cell(anno), _safe_csv_cell(r["month"]), _safe_csv_cell(r["week"]),
                        _safe_csv_cell(r["audio"]), _safe_csv_cell(r["video"]),
                        _safe_csv_cell(r["sabato"]), _safe_csv_cell(r["busy"])])
        return buf.getvalue()


# ══════════════════════════════════════════════════════════════
#  APPLICAZIONE
# ══════════════════════════════════════════════════════════════
class TurniApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Gestione Turni")
        self.geometry("880x700")
        self.minsize(760, 580)
        self.configure(bg=BG)
        self.resizable(True, True)
        configure_ttk_style()

        # Dati sessione
        self.operatori       = []
        self.operatori_norm  = []
        # [v3 BUG] anno corrente dinamico invece di hardcoded "2025"
        self.anno            = tk.StringVar(value=str(datetime.date.today().year))
        self.mesi            = []
        self.weeks_data      = []
        # [v4 UX] timeout configurabile dall'utente
        self.solver_timeout  = tk.StringVar(value=str(int(SOLVER_TIMEOUT)))

        # [CRITICO] inizializzazione garantita
        self._last_result_text = ""
        self._last_solver      = None

        # [CRITICO] controllo stato calcolo
        self._solving       = False
        self._cancel_event  = threading.Event()
        self._pending_close   = False
        self._pending_restart = False

        # Stato wizard
        self.current_step = 0
        self.steps        = ["Operatori & Anno", "Settimane", "Risultati"]

        # Per evitare ricostruzione inutile di step1
        self._step1_built  = False
        self._prev_ops:  list[str] = []
        self._prev_mesi: list[str] = []
        self._week_rows: list[dict]= []
        self._month_frames: dict[str, tk.Frame] = {}   # [v9 ARCH] init in __init__
        self._btn_save_session = None
        self._btn_load_session = None
        self._btn_modify = None
        self._btn_save_txt = None
        self._btn_save_csv = None
        self._btn_restart = None
        self._cancel_btn = None

        self._build_ui()
        self._show_step(0)
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ─── Layout ───────────────────────────────────────────────
    def _build_ui(self):
        self._build_header()
        self._build_stepbar()
        self._build_content()

    def _build_header(self):
        hdr = tk.Frame(self, bg=HEADER_BG, height=56)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="  Gestione Turni",
                 bg=HEADER_BG, fg="white",
                 font=FONT_TITLE).pack(side="left", padx=12, pady=10)
        mb = tk.Frame(hdr, bg=HEADER_BG)
        mb.pack(side="right", padx=12)
        self._btn_save_session = styled_btn(mb, "Salva sessione",  self._save_session,
                                            bg=HEADER_BTN)
        self._btn_save_session.pack(side="left", padx=4)
        self._btn_load_session = styled_btn(mb, "Carica sessione", self._load_session,
                                            bg=HEADER_BTN)
        self._btn_load_session.pack(side="left", padx=4)
        if not ORTOOLS_OK:
            tk.Label(hdr, text="ATTENZIONE: ortools mancante!",
                     bg=HEADER_BG, fg=INFO_WARN,
                     font=FONT_SMALL).pack(side="right", padx=8)

    def _build_stepbar(self):
        bar = tk.Frame(self, bg=PANEL_BG, height=46,
                       highlightbackground=BORDER, highlightthickness=1)
        bar.pack(fill="x")
        bar.pack_propagate(False)
        self._step_labels = []
        for i, name in enumerate(self.steps):
            lbl = tk.Label(bar, text=f"  {i+1}. {name}  ",
                           bg=PANEL_BG, fg=MUTED, font=FONT_SMALL,
                           padx=8, pady=12)
            lbl.pack(side="left")
            self._step_labels.append(lbl)
            if i < len(self.steps)-1:
                tk.Label(bar, text=">", bg=PANEL_BG,
                         fg=BORDER, font=(_FONT_FAMILY, 14)).pack(side="left")

    def _build_content(self):
        self.content = tk.Frame(self, bg=BG)
        self.content.pack(fill="both", expand=True)
        self.frames = {}
        for i in range(len(self.steps)):
            f = tk.Frame(self.content, bg=BG)
            f.place(relx=0, rely=0, relwidth=1, relheight=1)
            self.frames[i] = f
        self._build_step0(self.frames[0])
        self._build_step1(self.frames[1])
        self._build_step2(self.frames[2])

    def _show_step(self, n):
        self.current_step = n
        for i, lbl in enumerate(self._step_labels):
            if i == n:
                lbl.config(fg=ACCENT, font=FONT_BOLD, bg=MONTH_HDR_BG)
            elif i < n:
                lbl.config(fg=SUCCESS, font=FONT_SMALL, bg=PANEL_BG)
            else:
                lbl.config(fg=MUTED, font=FONT_SMALL, bg=PANEL_BG)
        self.frames[n].lift()

    def _set_solver_ui_state(self, solving: bool):
        state = "disabled" if solving else "normal"
        for btn in (
            self._btn_save_session,
            self._btn_load_session,
            self._btn_modify,
            self._btn_save_txt,
            self._btn_save_csv,
            self._btn_restart,
        ):
            if btn is not None:
                btn.config(state=state)
        if self._cancel_btn is not None:
            self._cancel_btn.config(state="normal" if solving else "disabled")

    def _request_cancel(self):
        if not self._solving:
            return
        self._cancel_event.set()
        cp_solver = getattr(self._last_solver, "cp_solver", None)
        for method_name in ("stop_search", "StopSearch"):
            stop_method = getattr(cp_solver, method_name, None)
            if callable(stop_method):
                try:
                    stop_method()
                except Exception:
                    logger.exception("Impossibile interrompere il solver con %s", method_name)
                break
        if self._cancel_btn is not None:
            self._cancel_btn.config(state="disabled")

    def _perform_restart(self):
        self.weeks_data        = []
        self._last_solver      = None
        self._last_result_text = ""
        self._pending_restart  = False
        self._result_subtitle.config(text="", fg=MUTED)
        self._set_result("")
        if self.mesi:
            self._build_step1_content(add_default_row=False)
            self._step1_built = True
        else:
            self._step1_built = False
            self._week_rows = []
            self._month_frames = {}
        self._show_step(0)

    def _on_close(self):
        """[CRITICO] chiusura sicura anche durante il calcolo."""
        if self._solving:
            if not messagebox.askyesno(
                    "Calcolo in corso",
                    "Il calcolo e' ancora in esecuzione.\n"
                    "Vuoi annullarlo e uscire?", parent=self):
                return
            self._pending_close = True
            self._request_cancel()
            self._result_subtitle.config(text="Annullamento in corso...", fg=WARN)
            self._set_result("Annullamento in corso. La finestra verra' chiusa al termine del solver.\n")
            self._set_solver_ui_state(True)
            if self._cancel_btn is not None:
                self._cancel_btn.config(state="disabled")
            return
        self.destroy()

    # ═══ STEP 0 – Operatori & Anno ════════════════════════════
    def _build_step0(self, parent):
        wrap = tk.Frame(parent, bg=BG)
        wrap.pack(expand=True)

        tk.Label(wrap, text="Operatori e Anno",
                 bg=BG, fg=TEXT, font=FONT_TITLE).pack(pady=(28, 4))
        tk.Label(wrap, text="Inserisci i dati base per pianificare i turni.",
                 bg=BG, fg=MUTED, font=FONT_LABEL).pack(pady=(0, 16))

        c = card(wrap)
        c.pack(padx=40, pady=8, fill="x")
        inner = tk.Frame(c, bg=PANEL_BG)
        inner.pack(padx=24, pady=20, fill="x")
        inner.columnconfigure(1, weight=1)

        def lbl(row, text):
            tk.Label(inner, text=text, bg=PANEL_BG,
                     fg=TEXT, font=FONT_BOLD).grid(
                row=row, column=0, sticky="nw", pady=8)

        def entry(row, var=None, width=44):
            e = tk.Entry(inner, font=FONT_LABEL, width=width,
                         relief="solid",
                         highlightbackground=BORDER, highlightthickness=1,
                         textvariable=var)
            e.grid(row=row, column=1, sticky="ew", padx=12, pady=8)
            return e

        lbl(0, "Anno")
        self._anno_entry = entry(0, var=self.anno, width=12)

        lbl(1, "Mesi (separati da virgola)")
        self._mesi_entry = entry(1)
        self._mesi_entry.insert(0, "Gennaio, Febbraio")

        lbl(2, "Operatori\n(uno per riga o virgola)")
        tf = tk.Frame(inner, bg=PANEL_BG)
        tf.grid(row=2, column=1, sticky="ew", padx=12, pady=8)
        self._ops_text = tk.Text(tf, font=FONT_LABEL, width=44, height=5,
                                  relief="solid", wrap="word",
                                  highlightbackground=BORDER,
                                  highlightthickness=1)
        self._ops_text.pack(side="left", fill="both", expand=True)
        # [v3 BUG] scrollbar collegata al Text widget (yscrollcommand mancante in v2)
        _ops_sb = ttk.Scrollbar(tf, command=self._ops_text.yview)
        self._ops_text.configure(yscrollcommand=_ops_sb.set)
        _ops_sb.pack(side="right", fill="y")
        self._ops_text.insert("1.0",
            "Mario Rossi\nLuca Bianchi\nAnna Verdi\nGiulia Neri")

        # [v4 UX] Timeout solver configurabile
        lbl(3, "Timeout solver (sec)")
        timeout_frame = tk.Frame(inner, bg=PANEL_BG)
        timeout_frame.grid(row=3, column=1, sticky="w", padx=12, pady=8)
        tk.Entry(timeout_frame, textvariable=self.solver_timeout,
                 font=FONT_LABEL, width=8, relief="solid",
                 highlightbackground=BORDER, highlightthickness=1).pack(side="left")
        tk.Label(timeout_frame,
                 text="  (aumenta se il solver non trova soluzioni in tempo)",
                 bg=PANEL_BG, fg=MUTED, font=FONT_SMALL).pack(side="left")

        styled_btn(wrap, "Avanti  >", self._step0_next).pack(pady=20)

    def _step0_next(self):
        # [v9 BUG] guardia _solving: navigare manualmente allo step 0 mentre il
        # solver gira e cliccare "Avanti" ricostruisce step1 con thread vivo
        if self._solving:
            messagebox.showwarning("Calcolo in corso",
                "Attendi il termine del calcolo prima di modificare i dati.",
                parent=self)
            return

        try:
            anno, timeout_str, new_ops, new_norm, new_mesi, dups, dup_mesi = _parse_step0_inputs(
                self.anno.get(),
                self._mesi_entry.get(),
                self._ops_text.get("1.0", "end"),
                self.solver_timeout.get(),
            )
        except SessionValidationError as e:
            messagebox.showerror("Errore", str(e), parent=self)
            return

        self.anno.set(anno)
        self.solver_timeout.set(timeout_str)

        if dups:
            messagebox.showwarning("Attenzione",
                "Operatori duplicati rimossi:\n" + "\n".join(dups),
                parent=self)

        if dup_mesi:
            messagebox.showwarning("Attenzione",
                "Mesi duplicati rimossi:\n" + "\n".join(dup_mesi),
                parent=self)

        ops_ok  = new_ops  == self._prev_ops
        mesi_ok = new_mesi == self._prev_mesi

        # [UX] Se nulla e' cambiato, vai direttamente a step1
        if self._step1_built and ops_ok and mesi_ok:
            self.operatori = new_ops
            self.operatori_norm = new_norm
            self.mesi = new_mesi
            self._show_step(1)
            return

        # Se ci sono dati e qualcosa e' cambiato, avvisa
        if self._step1_built and self._week_rows and not (ops_ok and mesi_ok):
            if not messagebox.askyesno(
                    "Attenzione",
                    "Hai modificato operatori o mesi.\n"
                    "Le settimane inserite verranno azzerate. Continuare?",
                    parent=self):
                return

        self.operatori      = new_ops
        self.operatori_norm = new_norm
        self.mesi           = new_mesi
        self._prev_ops      = list(new_ops)
        self._prev_mesi     = list(new_mesi)
        self.weeks_data     = []

        self._build_step1_content()
        self._step1_built = True
        self._show_step(1)

    # ═══ STEP 1 – Settimane ═══════════════════════════════════
    def _build_step1(self, parent):
        self._step1_parent = parent

    def _build_step1_content(self, add_default_row: bool = True):
        for w in self._step1_parent.winfo_children():
            w.destroy()
        self._week_rows    = []
        self._month_frames = {}

        tk.Label(self._step1_parent, text="Pianificazione Settimane",
                 bg=BG, fg=TEXT, font=FONT_TITLE).pack(pady=(20, 2))
        tk.Label(self._step1_parent,
                 text="Aggiungi settimane e spunta gli operatori non disponibili.",
                 bg=BG, fg=MUTED, font=FONT_LABEL).pack(pady=(0, 8))

        sf = ScrollableFrame(self._step1_parent)
        sf.pack(fill="both", expand=True, padx=20)

        for mese in self.mesi:
            mc = card(sf.inner)
            mc.pack(fill="x", pady=6, padx=4)
            hdr = tk.Frame(mc, bg=MONTH_HDR_BG)
            hdr.pack(fill="x")
            tk.Label(hdr, text=f"  {mese}",
                     bg=MONTH_HDR_BG, fg=ACCENT,
                     font=FONT_BOLD, anchor="w").pack(
                side="left", padx=10, pady=8)
            styled_btn(hdr, "+ Aggiungi settimana",
                       lambda m=mese: self._add_week_row(m),
                       bg=ACCENT_LIGHT, fg=ACCENT).pack(
                side="right", padx=10, pady=6)
            rows_frame = tk.Frame(mc, bg=PANEL_BG)
            rows_frame.pack(fill="x", padx=8, pady=4)
            self._month_frames[mese] = rows_frame
            # [v5 ARCH] riga di default opzionale: evita costruisci-poi-distruggi
            if add_default_row:
                self._add_week_row(mese)

        nav = tk.Frame(self._step1_parent, bg=BG)
        nav.pack(pady=14)
        styled_btn(nav, "<  Indietro",
                   lambda: self._show_step(0),
                   bg=BTN_SECONDARY, fg=TEXT).pack(side="left", padx=8)
        styled_btn(nav, "Genera Turni  >",
                   self._step1_next, bg=SUCCESS).pack(side="left", padx=8)

    def _add_week_row(self, mese):
        rows_frame = self._month_frames[mese]
        # [v6 ARCH] sum() più chiaro della list comprehension, stessa complessità
        row_idx    = sum(1 for r in self._week_rows if r["month"] == mese)

        row = tk.Frame(rows_frame, bg=PANEL_BG,
                       highlightbackground=BTN_SECONDARY, highlightthickness=1)
        row.pack(fill="x", pady=3, padx=4)

        lf = tk.Frame(row, bg=PANEL_BG)
        lf.pack(side="left", padx=8, pady=6)
        # [UX] salviamo riferimento per rinumerare dopo eliminazione
        num_lbl = tk.Label(lf, text=f"Settimana {row_idx+1}:",
                           bg=PANEL_BG, fg=TEXT, font=FONT_BOLD)
        num_lbl.pack(anchor="w")
        name_var = tk.StringVar(value=f"Sett.{row_idx+1} {mese[:3]}")
        tk.Entry(lf, textvariable=name_var, font=FONT_LABEL,
                 width=22, relief="solid",
                 highlightbackground=BORDER, highlightthickness=1).pack(pady=2)

        bf = tk.Frame(row, bg=PANEL_BG)
        bf.pack(side="left", padx=16, pady=6, fill="x", expand=True)
        tk.Label(bf, text="Operatori busy:",
                 bg=PANEL_BG, fg=MUTED, font=FONT_SMALL).pack(anchor="w")
        chk = tk.Frame(bf, bg=PANEL_BG)
        chk.pack(anchor="w")
        busy_vars = {}
        cols = 4
        for idx, op in enumerate(self.operatori):
            var = tk.BooleanVar(value=False)
            busy_vars[idx] = var
            tk.Checkbutton(chk, text=op, variable=var,
                           bg=PANEL_BG, fg=TEXT, font=FONT_SMALL,
                           activebackground=PANEL_BG,
                           selectcolor=ACCENT_LIGHT).grid(
                row=idx//cols, column=idx%cols,
                sticky="w", padx=8, pady=1)

        rd = {
            "month":      mese,
            "name_var":   name_var,
            "busy_vars":  busy_vars,
            "row_widget": row,
            "num_label":  num_lbl,
        }
        self._week_rows.append(rd)
        styled_btn(row, "X",
                   lambda r=rd, rw=row: self._remove_week_row(r, rw),
                   bg=DANGER_LIGHT, fg=DANGER).pack(side="right", padx=8, pady=6)

    def _remove_week_row(self, rd, rw):
        mese = rd["month"]
        rw.destroy()
        if rd in self._week_rows:
            self._week_rows.remove(rd)
        self._renumber_weeks(mese)  # [UX] rinumerazione

    def _renumber_weeks(self, mese):
        # [v8 ARCH] generatore: evita lista temporanea allocata a ogni chiamata
        for i, rd in enumerate(r for r in self._week_rows if r["month"] == mese):
            rd["num_label"].config(text=f"Settimana {i+1}:")

    def _step1_next(self):
        # [v3 BUG] previene doppio avvio del solver
        if self._solving:
            messagebox.showinfo("Calcolo in corso",
                "Il calcolo e' ancora in esecuzione.", parent=self)
            return
        if not self._week_rows:
            messagebox.showerror("Errore",
                "Aggiungi almeno una settimana.", parent=self)
            return

        raw_weeks = [
            {
                "month": rd["month"],
                "week": rd["name_var"].get(),
                "busy_indices": [i for i, v in rd["busy_vars"].items() if v.get()],
            }
            for rd in self._week_rows
        ]

        try:
            canonical_weeks = _validate_week_entries(
                raw_weeks,
                declared_months=self.mesi,
                error_prefix="Errore Step 1",
            )
            _validate_solver_ready_weeks(
                canonical_weeks,
                operator_count=len(self.operatori),
                error_prefix="Errore Step 1",
            )
        except SessionValidationError as e:
            msg = str(e)
            if msg.startswith("Errore Step 1: "):
                msg = msg[len("Errore Step 1: "):]
            messagebox.showerror("Errori", msg, parent=self)
            return

        new_weeks_data = []
        for week in canonical_weeks:
            busy_norm = [self.operatori_norm[i] for i in week["busy_indices"]]
            available = [
                i for i, op in enumerate(self.operatori_norm)
                if op not in busy_norm
            ]
            new_weeks_data.append({
                "month": week["month"],
                "week": week["week"],
                "busy": busy_norm,
                "available": available,
            })

        if not new_weeks_data:
            messagebox.showerror("Errore",
                "Nessuna settimana valida.", parent=self)
            return
        self.weeks_data = _order_weeks_by_declared_months(new_weeks_data, self.mesi)
        self._show_step(2)
        self._run_solver()

    # ═══ STEP 2 – Risultati ═══════════════════════════════════
    def _build_step2(self, parent):
        tk.Label(parent, text="Risultati",
                 bg=BG, fg=TEXT, font=FONT_TITLE).pack(pady=(20, 2))
        self._result_subtitle = tk.Label(parent, text="",
                                          bg=BG, fg=MUTED, font=FONT_LABEL)
        self._result_subtitle.pack(pady=(0, 6))

        rc = card(parent)
        rc.pack(fill="both", expand=True, padx=20, pady=4)
        self._result_text = tk.Text(rc, font=FONT_MONO, wrap="none",
                                     bg=PANEL_BG, fg=TEXT, relief="flat",
                                     state="disabled", padx=12, pady=8)
        sb_y = ttk.Scrollbar(rc, command=self._result_text.yview)
        sb_x = ttk.Scrollbar(rc, orient="horizontal",
                              command=self._result_text.xview)
        self._result_text.configure(yscrollcommand=sb_y.set,
                                     xscrollcommand=sb_x.set)
        sb_y.pack(side="right", fill="y")
        sb_x.pack(side="bottom", fill="x")
        self._result_text.pack(fill="both", expand=True)

        # Barra progresso + pulsante annulla [UX fix #5]
        self._prog_frame = tk.Frame(parent, bg=BG)
        self._prog_frame.pack(pady=4)
        self._progress = ttk.Progressbar(
            self._prog_frame, mode="indeterminate", length=280)
        self._progress.pack(side="left", padx=6)
        self._cancel_btn = styled_btn(
            self._prog_frame, "Annulla", self._cancel_solving, bg=DANGER)
        self._cancel_btn.pack(side="left", padx=6)
        self._prog_frame.pack_forget()

        nav = tk.Frame(parent, bg=BG)
        nav.pack(pady=10)
        self._btn_modify = styled_btn(nav, "<  Modifica",
                                      lambda: self._show_step(1),
                                      bg=BTN_SECONDARY, fg=TEXT)
        self._btn_modify.pack(side="left", padx=6)
        self._btn_save_txt = styled_btn(nav, "Salva TXT",
                                        lambda: self._export("txt"),
                                        bg=SUCCESS)
        self._btn_save_txt.pack(side="left", padx=6)
        self._btn_save_csv = styled_btn(nav, "Salva CSV",  # [EXTRA]
                                        lambda: self._export("csv"),
                                        bg=CYAN)
        self._btn_save_csv.pack(side="left", padx=6)
        self._btn_save_docx = styled_btn(nav, "Salva DOCX",
                                         lambda: self._export("docx"),
                                         bg="#2563EB")
        self._btn_save_docx.pack(side="left", padx=6)
        self._btn_restart = styled_btn(nav, "Nuovo calcolo",
                                       self._restart, bg=WARN)
        self._btn_restart.pack(side="left", padx=6)

    def _run_solver(self):
        """[CRITICO] snapshot immutabili passati al thread."""
        ops_snap = list(self.operatori)
        # [v5 BUG] deep copy delle liste interne: shallow dict() lasciava
        # "available" e "busy" condivisi per riferimento col thread principale
        weeks_snap = [
            {**w, "available": list(w["available"]), "busy": list(w["busy"])}
            for w in self.weeks_data
        ]

        self._solving      = True
        self._cancel_event = threading.Event()
        self._pending_close = False
        self._pending_restart = False
        self._last_solver  = TurniSolver(ops_snap, weeks_snap)

        # [v6 BUG] reset colore subtitle: senza fg=MUTED il colore del run
        # precedente (verde/arancio/rosso) restava durante il calcolo in corso
        self._result_subtitle.config(text="Ottimizzazione in corso...", fg=MUTED)
        self._prog_frame.pack(pady=4)
        self._progress.start(10)
        self._set_solver_ui_state(True)
        self._set_result("Calcolo in corso, attendere...\n")

        ev         = self._cancel_event
        # [v8 BUG] float() era fuori dal worker e non protetto:
        # un timeout non valido (es. caricato da sessione corrotta) causava ValueError
        # che lasciava _solving=True permanente. Ora fallback esplicito a SOLVER_TIMEOUT.
        try:
            timeout = float(self.solver_timeout.get())
            if timeout <= 0:
                raise ValueError("timeout non positivo")
        except ValueError:
            timeout = SOLVER_TIMEOUT
            self.solver_timeout.set(str(int(SOLVER_TIMEOUT)))
            logger.warning("Timeout non valido in _run_solver, usato default %.1f", SOLVER_TIMEOUT)
        # [v7 BUG] snapshot locale: il worker non legge self._last_solver a runtime
        # (protezione contro future rimozioni della guardia _solving)
        solver_ref = self._last_solver

        def worker():
            # [v4 BUG] try/except: _solving non resta bloccato su eccezione imprevista
            try:
                txt = solver_ref.solve(ev, timeout=timeout)
            except Exception as exc:
                logger.exception("Errore imprevisto nel solver")
                txt = f"Errore imprevisto durante il calcolo:\n{exc}"
            try:
                self.after(0, lambda: self._display_result(txt))
            except tk.TclError:
                logger.debug("UI gia' chiusa: risultato solver ignorato")

        threading.Thread(target=worker, daemon=True).start()

    def _cancel_solving(self):
        self._request_cancel()
        self._result_subtitle.config(text="Annullamento in corso...", fg=WARN)
        self._set_result("Annullamento in corso, attendere...\n")

    def _display_result(self, text):
        # [v3 BUG] la finestra potrebbe essere stata chiusa durante il calcolo
        if not self.winfo_exists():
            return
        self._solving = False
        self._progress.stop()
        self._prog_frame.pack_forget()
        self._set_solver_ui_state(False)
        self._last_result_text = text
        # [v3 ARCH] usa solved_ok esplicito invece di text.startswith("-")
        ok = self._last_solver is not None and self._last_solver.solved_ok
        cancelled = self._last_solver is not None and self._last_solver.cancelled
        partial = (self._last_solver is not None
                   and self._last_solver.partial_result_available)
        # [v5 UX] colore subtitle differenziato: verde=ok, arancio=annullato, rosso=errore
        if ok:
            subtitle, color = "Pianificazione completata.", SUCCESS
        elif cancelled and partial:
            subtitle, color = "Calcolo annullato: mostrata l'ultima soluzione parziale trovata.", WARN
        elif cancelled:
            subtitle, color = "Calcolo annullato.", WARN
        else:
            subtitle, color = text.split("\n")[0], DANGER
        self._result_subtitle.config(text=subtitle, fg=color)
        self._set_result(text)
        if self._pending_close:
            self._pending_close = False
            self.destroy()
            return
        if self._pending_restart:
            self._perform_restart()

    def _set_result(self, text):
        self._result_text.config(state="normal")
        self._result_text.delete("1.0", "end")
        self._result_text.insert("1.0", text)
        self._result_text.config(state="disabled")

    def _restart(self):
        if self._solving:
            self._pending_restart = True
            self._request_cancel()
            self._result_subtitle.config(
                text="Annullamento in corso... verra' avviato un nuovo calcolo.",
                fg=WARN,
            )
            self._set_result(
                "Annullamento in corso. Il reset verra' eseguito al termine del solver.\n")
            return
        self._perform_restart()

    # ─── Export ──────────────────────────────────────────────
    def _export(self, fmt):
        # [v3 ARCH] controllo coerente con _display_result tramite solved_ok
        solver_ok = self._last_solver is not None and self._last_solver.solved_ok
        if not solver_ok:
            messagebox.showinfo("Info",
                "Nessun risultato da esportare.", parent=self)
            return
        anno = self.anno.get()
        path = None
        if fmt == "txt":
            path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Testo", "*.txt"), ("Tutti", "*.*")],
                initialfile=f"turni_{anno}.txt", parent=self)
            if path:
                try:
                    _write_text_file_atomic(path, self._last_result_text)
                except OSError as e:
                    messagebox.showerror("Errore salvataggio",
                        f"Impossibile scrivere il file:\n{e}", parent=self)
                    return
        elif fmt == "csv":
            path = filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV", "*.csv"), ("Tutti", "*.*")],
                initialfile=f"turni_{anno}.csv", parent=self)
            if path:
                try:
                    _write_text_file_atomic(
                        path,
                        self._last_solver.format_csv(anno),
                        newline="",
                    )
                except OSError as e:
                    messagebox.showerror("Errore salvataggio",
                        f"Impossibile scrivere il file:\n{e}", parent=self)
                    return
        elif fmt == "docx":
            path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx"), ("Tutti", "*.*")],
                initialfile=f"turni_{anno}.docx", parent=self)
            if path:
                try:
                    build_turni_docx(
                        path,
                        anno,
                        list(self._last_solver.result_rows),
                        list(self._last_solver.operatori),
                        list(self._last_solver.counts),
                        self._last_solver.status_str,
                        self._last_solver.diff_val,
                        self._last_solver.penalty_val,
                    )
                except OSError as e:
                    messagebox.showerror("Errore salvataggio",
                        f"Impossibile scrivere il file:\n{e}", parent=self)
                    return
                except Exception as e:
                    logger.exception("Errore durante la creazione del DOCX")
                    messagebox.showerror("Errore esportazione DOCX",
                        f"Impossibile creare il documento Word:\n{e}", parent=self)
                    return
        if path:
            messagebox.showinfo("Salvato",
                f"File salvato:\n{path}", parent=self)

    # ─── Sessione JSON [ARCH] ────────────────────────────────
    def _save_session(self):
        # [v4 BUG] impedisce salvataggio di stato inconsistente durante il calcolo
        if self._solving:
            messagebox.showwarning("Calcolo in corso",
                "Attendi il termine del calcolo prima di salvare la sessione.",
                parent=self)
            return

        try:
            anno, timeout_str, parsed_ops, _parsed_norms, parsed_mesi, dup_ops, dup_mesi = _parse_step0_inputs(
                self.anno.get(),
                self._mesi_entry.get(),
                self._ops_text.get("1.0", "end"),
                self.solver_timeout.get(),
            )
        except SessionValidationError as e:
            messagebox.showerror("Errore",
                f"Impossibile salvare la sessione:\n{e}", parent=self)
            return

        if dup_ops:
            messagebox.showwarning("Attenzione",
                "Operatori duplicati rimossi nel salvataggio:\n" + "\n".join(dup_ops),
                parent=self)
        if dup_mesi:
            messagebox.showwarning("Attenzione",
                "Mesi duplicati rimossi nel salvataggio:\n" + "\n".join(dup_mesi),
                parent=self)

        if self._step1_built and (parsed_ops != self.operatori or parsed_mesi != self.mesi):
            messagebox.showerror("Errore",
                "Hai modificato operatori o mesi nello Step 0 ma non hai ancora applicato "
                "le modifiche. Clicca 'Avanti' per confermarle prima di salvare la sessione.",
                parent=self)
            return

        raw_session = {
            "anno": anno,
            "operatori": parsed_ops,
            "mesi": parsed_mesi,
            "solver_timeout": timeout_str,
            "weeks": [
                {
                    "month": rd["month"],
                    "week": rd["name_var"].get(),
                    "busy_indices": [i for i, v in rd["busy_vars"].items() if v.get()],
                }
                for rd in self._week_rows
            ],
        }

        try:
            session = _validate_session_payload(raw_session)
            session["solver_timeout"] = timeout_str
        except SessionValidationError as e:
            messagebox.showerror("Errore",
                f"Impossibile salvare la sessione:\n{e}", parent=self)
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("Sessione", "*.json"), ("Tutti", "*.*")],
            initialfile=f"sessione_{session['anno']}.json",
            parent=self)
        if path:
            try:
                _write_text_file_atomic(
                    path,
                    json.dumps(session, ensure_ascii=False, indent=2),
                )
            except OSError as e:
                messagebox.showerror("Errore salvataggio",
                    f"Impossibile scrivere il file:\n{e}", parent=self)
                return
            messagebox.showinfo("Salvato",
                f"Sessione salvata:\n{path}", parent=self)

    def _load_session(self):
        # [v5 BUG] guard: caricamento durante calcolo produce stato incoerente
        if self._solving:
            messagebox.showwarning("Calcolo in corso",
                "Attendi il termine del calcolo prima di caricare una sessione.",
                parent=self)
            return
        path = filedialog.askopenfilename(
            filetypes=[("Sessione", "*.json"), ("Tutti", "*.*")],
            parent=self)
        if not path:
            return
        try:
            with open(path, encoding="utf-8") as f:
                raw_session = json.load(f)
        except Exception as e:
            messagebox.showerror("Errore",
                f"Impossibile leggere il file:\n{e}", parent=self)
            return

        try:
            session = _validate_session_payload(raw_session)
        except SessionValidationError as e:
            messagebox.showerror("Errore", str(e), parent=self)
            return

        ops = session["operatori"]
        mesi = session["mesi"]
        weeks = session["weeks"]
        anno_val = session["anno"]

        # [v6 BUG] reset risultati precedenti: evita export di dati di un'altra sessione
        self._last_solver      = None
        self._last_result_text = ""

        # Ripristina step0
        self.anno.set(anno_val)
        if "solver_timeout" in session:
            self.solver_timeout.set(session["solver_timeout"])
        self._ops_text.delete("1.0", "end")
        self._ops_text.insert("1.0", "\n".join(ops))
        self._mesi_entry.delete(0, "end")
        self._mesi_entry.insert(0, ", ".join(mesi))

        self.operatori      = ops
        self.operatori_norm = [normalize_name(o) for o in ops]
        self.mesi           = mesi
        self._prev_ops      = list(self.operatori)
        self._prev_mesi     = list(self.mesi)

        # [v5 ARCH] add_default_row=False: evita il ciclo costruisci-poi-distruggi
        self._build_step1_content(add_default_row=False)
        self._step1_built = True

        # Ripristina settimane
        by_month = _group_weeks_by_normalized_month(weeks)

        for mese in self.mesi:
            for w in by_month.get(normalize_name(mese), []):
                self._add_week_row(mese)
                # [v9 ARCH] O(1): la riga appena aggiunta è sempre l'ultima
                rd = self._week_rows[-1]
                rd["name_var"].set(w.get("week", ""))
                for idx in w.get("busy_indices", []):
                    if idx in rd["busy_vars"]:
                        rd["busy_vars"][idx].set(True)
            self._renumber_weeks(mese)

        self._show_step(1)
        messagebox.showinfo("Caricato",
            "Sessione caricata. Verifica i dati e clicca 'Genera Turni'.",
            parent=self)


# ══════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    # [v4 ARCH] logging configurato al livello WARNING di default
    logging.basicConfig(
        level=logging.WARNING,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )
    if not ORTOOLS_OK:
        _r = tk.Tk()
        _r.withdraw()
        messagebox.showwarning(
            "Dipendenza mancante",
            "ortools non e' installato.\n\n"
            "Installa con:\n    pip install ortools")
        _r.destroy()

    TurniApp().mainloop()
